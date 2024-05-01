#!/usr/bin/python3
# Ejecutar en la carpeta raiz (en el nivel de las carpetas EXTERNO/INTERNO )
import os
from openpyxl.styles import PatternFill, Border, Side, Alignment, Protection, Font
from openpyxl.chart.shapes import GraphicalProperties
from openpyxl.chart.marker import DataPoint
from openpyxl.chart.label import DataLabelList
from lxml import etree
import subprocess
import argparse
import openpyxl
import math
import pandas as pd
import sys
from openpyxl.styles import PatternFill, Border, Side, Alignment, Protection, Font
from openpyxl.chart import PieChart3D, PieChart, ProjectedPieChart, BarChart, Reference
from openpyxl.chart.shapes import GraphicalProperties
from openpyxl.chart.marker import DataPoint
from openpyxl.chart.label import DataLabelList
import sqlite3
import pprint
import time
from mergeDB import merge_databases
from PIL import Image, ImageDraw, ImageFont 
import argparse
import json
from openpyxl.cell.cell import ILLEGAL_CHARACTERS_RE
import re 
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml.shared import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

############# func retest ###

def create_custom_table(doc, title, index):
    """
    Creates a custom table in the given Word document.
    """
    # Add a table with 6 rows and 2 columns
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    # title
    set_cell_text(table.cell(0, 0), str(index), "Impact", 36, True, False)
   # set_cell_text(table.cell(0, 1), title, "Impact", 22, True, False)
    set_cell_text(table.cell(0, 1).merge(table.cell(0, 2)), title, "Impact", 22, True, False)

    #estado
    set_cell_text(table.cell(1, 0), "ESTADO", None, None, True, False)
    table.cell(1, 1).merge(table.cell(1, 2))
    
    # Configure merged rows with specific backgrounds and make them bold as needed
    merge_and_format_cells(table, 2, "EVIDENCIA", "D3D3D3", True)
    merge_and_format_cells(table, 3, "", "FFFFFF", False)
    merge_and_format_cells(table, 4, "CONCLUSIÓN", "D3D3D3", True)
    merge_and_format_cells(table, 5, "", "FFFFFF", False)

# Función para establecer el color de fondo de una celda
def set_cell_background(cell, fill):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill))
    cell._element.get_or_add_tcPr().append(shading_elm)


def set_cell_text(cell, text, font_name, font_size, center_align, is_bold):
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = paragraph.add_run(text)
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    run.font.bold = is_bold
    if center_align:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def merge_and_format_cells(table, row_index, text, bg_color, is_bold):
    merged_cell = table.cell(row_index, 0).merge(table.cell(row_index, 2))
    set_cell_background_color(merged_cell, bg_color)
    set_cell_text(merged_cell, text, None, None, True, is_bold)

def set_cell_background_color(cell, fill):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

#############
parser = argparse.ArgumentParser()
parser.add_argument("--phishing", "-p", help="phishing report")
parser.add_argument("--empresa", "-e", help="phishing report") # ["RADICAL","RORIK"]
parser.parse_args()
args = parser.parse_args()
empresa = args.empresa


# Add long and short argument
matrizRecomendaciones_INTERNO = []
matrizRecomendaciones_EXTERNO = []
mainDB_INTERNO = [] 

# Definir los archivos para los informes
#empresas = ["RADICAL","RORIK"]
wbEjecutivo = openpyxl.Workbook() # el informe ejecutivo es igual para todos
wbMatrizVul = openpyxl.Workbook() # Matriz de vulnerabilidades es igual para todos

# Fuentes
Font_MuyAlto = Font(name='Century Gothic', size=12, bold=True, color='FC611C')
Font_Alto = Font(name='Century Gothic', size=12, bold=True, color='FFC000')
Font_Moderado = Font(name='Century Gothic', size=12, bold=True, color='FFFF00')
Font_Bajo = Font(name='Century Gothic', size=12, bold=True, color='00B222')
Font_Info = Font(name='Century Gothic', size=12, bold=True, color='026EBA')


Font_MuyAlto_RORIK = Font(name='Century Gothic', size=10, bold=True, color='FC611C')
Font_Alto_RORIK = Font(name='Century Gothic', size=10, bold=True, color='E36C0A')
Font_Moderado_RORIK = Font(name='Century Gothic', size=10, bold=True, color='FFC000')
Font_Bajo_RORIK = Font(name='Century Gothic', size=10, bold=True, color='00B222')
Font_muyBajo_RORIK = Font(name='Century Gothic', size=10, bold=True, color='026EBA')

italic24Font = Font(size=24, italic=True, bold=True)
Arial10 = Font(name='Century Gothic', size=10)
Arial11Bold = Font(name='Century Gothic', size=11, bold=True)
Arial14Bold = Font(name='Century Gothic', size=14, bold=True)
Arial12BoldWhite = Font(name='Century Gothic', size=12, bold=True, color='FFFFFF')
Arial10Bold = Font(name='Century Gothic', size=10, bold=True)
Calibri10 = Font(name='Century Gothic', size=10)
Calibri10Bold = Font(name='Century Gothic', size=10, bold=True)
Calibri12Bold = Font(name='Century Gothic', size=12, bold=True)
Calibri12 = Font(name='Century Gothic', size=12)
Calibri10BoldWhite = Font(name='Century Gothic', size=10, bold=True, color='FFFFFF')

Calibri22Bold = Font(name='Century Gothic', size=22, bold=True)

Impact22 = Font(name='Impact', size=22)
Impact36 = Font(name='Impact', size=36)

# Borde
thin_border = Border(left=Side(style='thin'),
                     right=Side(style='thin'),
                     top=Side(style='thin'),
                     bottom=Side(style='thin'))

# Rellenos
greyFill = PatternFill(start_color='d9d9d9', end_color='d9d9d9', fill_type='solid')
verdeFill = PatternFill(start_color='548DD4', end_color='548DD4', fill_type='solid')

#### RADICAL #####
# Mapeo de niveles de criticidad a colores
nivel_criticidad_colores = {
    "CRÍTICO": "FF0000",  # ROJO
    "ALTO": "FFC000",     # NARANJA
    "MEDIO": "FFFF00",    # AMARILLO
    "BAJO": "00B050",     # VERDE
    "INFORMATIVO": "95B3D7",  # AZUL CLARO
}


CRÍTICOFill = PatternFill(start_color='FF0000', end_color='FF0000', fill_type='solid')
altoFill = PatternFill(start_color='FFC000', end_color='FFC000', fill_type='solid')
medioFill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
bajoFill = PatternFill(start_color='00B050', end_color='00B050', fill_type='solid')
brownFill = PatternFill(start_color='EEECE1', end_color='EEECE1', fill_type='solid')

#### RORIK #####
AltoFill2 = PatternFill(start_color='FFC000', end_color='FFC000', fill_type='solid')
ModeradoFill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
bajoFill2 = PatternFill(start_color='92D050', end_color='92D050', fill_type='solid')
informativaFill = PatternFill(start_color='95B3D7', end_color='399BE0', fill_type='solid')
negroFill = PatternFill(start_color='17150D', end_color='17150D', fill_type='solid')

############################################################

# Busca el codigo CVE y la descripcion de una vulnerabilidad
def buscarDescripcion(nombreVul):
    CVE = "N/A"
    for i in range(total_vulnerabilidades_xml):
        current_cod = root[i].find("codVul").text
        descripcion = root[i].find("nombre").text        
        if nombreVul in current_cod:
            CVE = root[i].find("CVE").text
            break
    if "webarchivos" == nombreVul: descripcion = "Navegación forzada para descubrir archivos comunes"
    if "passwordAdivinado1" == nombreVul: descripcion = "Contraseña fácil de adivinar"
    if "passwordAdivinado2" == nombreVul: descripcion = "Contraseña fácil de adivinar"
    if "webdirectorios" == nombreVul: descripcion = "Navegación forzada para descubrir archivos directorios"
    if "googlehacking0" == nombreVul: descripcion = "Google dork: site:DOMINIO inurl:add"
    if "googlehacking1" == nombreVul: descripcion = "Google dork: site:DOMINIO inurl:edit"
    if "googlehacking2" == nombreVul: descripcion = "Google dork: site:DOMINIO intitle:index.of"
    if "googlehacking3" == nombreVul: descripcion = "Google dork: site:DOMINIO filetype:sql"
    if "googlehacking4" == nombreVul: descripcion = 'Google dork: site:DOMINIO "access denied for user"'
    if "googlehacking5" == nombreVul: descripcion = 'Google dork: site:DOMINIO intitle:"curriculum vitae"'
    if "googlehacking6" == nombreVul: descripcion = "Google dork: site:DOMINIO passwords|contrasenas|login|contrasena filetype:txt"
    if "googlehacking11" == nombreVul: descripcion = "Google dork: site:trello.com passwords|contrasenas|login|contrasena intext:DOMINIO"
    if "googlehacking13" == nombreVul: descripcion = 'Google dork: site:DOMINIO "Undefined index"'
    if "googlehacking14" == nombreVul: descripcion = 'Google dork: site:DOMINIO inurl:storage'
    return CVE+";"+descripcion


def generar_texto_vulnerabilidades(total_vul_criticas, total_vul_altas, total_vul_medias, total_vul_bajas, total_vul_info):
    total_vulnerabilidades = total_vul_criticas + total_vul_altas + total_vul_medias + total_vul_bajas + total_vul_info

    # Calcula los porcentajes, evitando la división por cero
    porcentaje_criticas = (total_vul_criticas / total_vulnerabilidades * 100) if total_vulnerabilidades else 0
    porcentaje_altas = (total_vul_altas / total_vulnerabilidades * 100) if total_vulnerabilidades else 0
    porcentaje_medias = (total_vul_medias / total_vulnerabilidades * 100) if total_vulnerabilidades else 0
    porcentaje_bajas = (total_vul_bajas / total_vulnerabilidades * 100) if total_vulnerabilidades else 0
    porcentaje_info = (total_vul_info / total_vulnerabilidades * 100) if total_vulnerabilidades else 0

    texto = f"Se identificaron {total_vulnerabilidades} vulnerabilidades de las cuales "
    partes = []
    
    if total_vul_criticas > 0:
        partes.append(f"{total_vul_criticas} son críticas ({porcentaje_criticas:.0f}%)")
    if total_vul_altas > 0:
        partes.append(f"{total_vul_altas} son de alto riesgo ({porcentaje_altas:.0f}%)")
    if total_vul_medias > 0:
        partes.append(f"{total_vul_medias} corresponden a un nivel de riesgo medio ({porcentaje_medias:.0f}%)")
    if total_vul_bajas > 0:
        partes.append(f"{total_vul_bajas} pertenecen a un nivel de riesgo bajo ({porcentaje_bajas:.0f}%) del total")
    if total_vul_info > 0:
        partes.append(f"{total_vul_info} son informativas ({porcentaje_info:.0f}%)")

    texto += ", ".join(partes)
    texto += "."

    return texto


#show detailed hosts
def create_detalle_activo(ip, port):
    # Definir los puertos para HTTPS
    https_ports = ["443", "8443", "4443", "4433", "10443"]
    # Definir los puertos para HTTP
    http_ports = ["80", "81", "82", "8000", "8080", "8081", "8082", "8090"]

    # Check if port is a digit and categorize accordingly
    if not port.isdigit():
        detalle_activo = ip
    elif port in https_ports:
        detalle_activo = "https://" + ip + ":" + port
    elif port in http_ports:
        detalle_activo = "http://" + ip + ":" + port
    else:
        detalle_activo = ip + ":" + port

    return detalle_activo


# Convertir el log en una imagen para el reporte         
#ruta,testType,ip,port,vulType
def log2Image(logPath,imageName):
    #print(f'logPath {logPath} imageName {imageName}')
    # INTERNO/sistemas/logs/vulnerabilidades/192.168.2.247_25_openrelay.txt    
    max_lines_log = 50 #maximo num de lineas del log a convertir en imagen
    logFile = open(logPath)
    vuln_detalles = logFile.read()
    if (vuln_detalles == ''):
        print ("archivo vacio")
        vuln_detalles = "N/A"
    logFile.close()

    vuln_detalles = vuln_detalles.replace("\n\n", "\n") # eliminar saltos de linea excesivos
    vuln_detalles = vuln_detalles.replace("\r", "")  # eliminar     
    vuln_detalles = vuln_detalles.replace("\t", "    ")  # eliminar     
    lista_vuln_detalles = vuln_detalles.split('\n') 
    lineas_log = len(lista_vuln_detalles)
    if lineas_log > max_lines_log:
        lineas_log = max_lines_log

    vuln_detalles = ""    
    i = 0
    #print(f"lista_vuln_detalles {lista_vuln_detalles}")
    imageWidth = 0                    
    for line in lista_vuln_detalles:
        line = line[:120]
        if i < max_lines_log :
            vuln_detalles = vuln_detalles + line + "\n"
            #print (f'vuln_detalles {vuln_detalles}')
            i = i + 1
            if (len(line) > imageWidth) and (len(line) < 121): 
               imageWidth = len(line) # determinar cadena con mas longitud (para el ancho de la imagen)
    
    imageWidth = imageWidth * 8
    imageHeight = lineas_log * 15

    #print (f'imageWidth {imageWidth}')
    #print (f'imageHeight {imageHeight}')

    # convertir a imagen
    imagenVulnerabilidad = Image.new('RGB', (imageWidth, imageHeight), color=(255, 255, 255))                    
    canvas = ImageDraw.Draw(imagenVulnerabilidad)
    fnt = ImageFont.truetype("arial.ttf", 12)
    canvas.text((12, 12), vuln_detalles,font=fnt, fill='#000000')
    imagenVulnerabilidad.save(imageName)

os.system(f"rm -rf INTERNO/consolidado 2>/dev/null")

root = etree.parse("/usr/share/lanscanner/vulnerabilidades.xml").getroot()
total_vulnerabilidades_xml = len(root.getchildren())
#print (f"total_vulnerabilidades_xml en xml {total_vulnerabilidades_xml}")

# Resultados Externos
result = subprocess.run(['find', 'EXTERNO','-iname', '.resultados.db'], stdout=subprocess.PIPE)
resultados_externos = result.stdout.decode("utf-8")
#print(f"resultados_externos {resultados_externos}")
resultados_EXTERNO_list = resultados_externos.split('\n')
resultados_EXTERNO_list.pop() # eliminar ultimo elemento (vacio)

# Resultados Internos
result = subprocess.run(['find', 'INTERNO','-iname', '.resultados.db'], stdout=subprocess.PIPE)
resultados_internos = result.stdout.decode("utf-8")
#print(f'resultados_internos_raw {resultados_internos}')

#convertir a lista
resultados_INTERNO_list = resultados_internos.split('\n')
resultados_INTERNO_list.pop() # eliminar ultimo elemento (vacio)
#print (f'resultados_INTERNO_list {resultados_INTERNO_list}')

# Borrar informes anteriores
os.system("rm *.xlsx 2>/dev/null")
os.system("rm *.png 2>/dev/null")



################## CREAR INFORME ############################
#for empresa in empresas:
    
# TODAS LAS PRUEBAS
dominio = ''
segmento = ''
total_host_analizados = 0
total_host_vulnerabilidades = 0
total_host_uniq_vulnerabilidades = 0

######## #Vulnerabilidades por riesgo ########
total_vul_criticas = 0
total_vul_altas = 0
total_vul_medias = 0
total_vul_bajas = 0
total_vul_info = 0
######

#####################################
                                                                                                    
vuln_no_log = ['smbrelayShare','passwordHost','smbrelayShell','archivosDefecto','mailPass','ddos','phishing','captcha','contenidoNoRelacionado','googleRedirect','defaultPassword','openWebservice','malwareEjecutable','escalamientoPrivilegios','noTLS','softwareObsoleto','malwareWeb','passwordDefecto','802-1x','perdidaAutenticacion','ListadoDirectorios','xss','exfiltracionInformacion','googleCachedContent','nullsession-localUsers','nullsession-domainUsers','fugaInformacion','LFI','FPD','sqli','SourceMap','RedInterna','domainTakeOver','CS-10','CS-50','headerSegurosFaltantes','contenidoPrueba','auexterna-databases','command','validacionInsuficiente','ListadoDirectorioArchivosSensibles','compartidoSMB','passwordZKSoftware','hikvisio-nuclei','AtaqueWifi','laptopControlesInsuficientes']  
###### Vulnerabilidades por vectors
total_vuln_externas = 0
total_vuln_internas = 0
#######

total_vul_criticas_EXTERNO = 0
total_vul_altas_EXTERNO = 0
total_vul_medias_EXTERNO  = 0
total_vul_bajas_EXTERNO = 0
total_vul_info_EXTERNO = 0

total_vul_criticas_INTERNO = 0
total_vul_altas_INTERNO= 0
total_vul_medias_INTERNO  = 0
total_vul_bajas_INTERNO = 0
total_vul_info_INTERNO = 0


###### Vulnerabilidades por activos ####
aplicacionEscritorio = 0
aplicacionMovil   = 0 
aplicacionWeb   = 0   
atm = 0
estacionesTrabajo = 0    
wifi = 0
servidores = 0
baseDatos  = 0
telefoniaIP  = 0
sistemaVigilancia  = 0
dispositivosRed  = 0       
personal = 0
otros = 0 # Impresoras, lectores de huella
#############

resultados_EXTERNO_count = 0
resultados_INTERNO_count = 0

#### Vulnerabilidades por categoria de vulnerabilidad ####
passwordDebil = 0
faltaParches = 0
errorConfiguracion = 0
programacionInsegura = 0
####################
totalPruebas = 0

############# CONSOLIDAR ESCANEOS INTERNOS ########################
#realizamos una copia de la base principal
if (len(resultados_INTERNO_list)>1):    
    os.system(f"mkdir -p INTERNO/consolidado/.datos 2>/dev/null")
    os.system(f"mkdir -p INTERNO/consolidado/.datos_archived 2>/dev/null")
    os.system(f"mkdir -p INTERNO/consolidado/reportes_archived 2>/dev/null")
    #copiamos la primera DB que sera la principal    
    os.system(f"cp /usr/share/lanscanner/.resultados.db INTERNO/consolidado/.resultados.db")
    for resultado_ruta in resultados_INTERNO_list:
        resultado_ruta = resultado_ruta.replace(".resultados.db", "")
        if "consolidado" not in resultado_ruta: 
            print(f' resultado_ruta {resultado_ruta}')
            print (f'cp {resultado_ruta}reportes/* {resultado_ruta}reportes_archived/')
            os.system(f"mkdir {resultado_ruta}.datos_archived 2>/dev/null")    
            os.system(f"mkdir {resultado_ruta}reportes_archived 2>/dev/null")    
            os.system(f'cp {resultado_ruta}.datos/* {resultado_ruta}.datos_archived/ 2>/dev/null')
            os.system(f'cp {resultado_ruta}reportes/* {resultado_ruta}reportes_archived/ 2>/dev/null')

            os.system(f"cp -r {resultado_ruta}* INTERNO/consolidado/ ")
            os.system(f"mkdir INTERNO/consolidado/.vulnerabilidades2/ 2>/dev/null")
            os.system(f"cp -r {resultado_ruta}.vulnerabilidades2/* INTERNO/consolidado/.vulnerabilidades2/ 2>/dev/null")
            os.system(f"cat {resultado_ruta}.datos_archived/total-host-vivos.txt >> INTERNO/consolidado/.datos_archived/total-host-vivos2.txt")
            os.system(f"cat {resultado_ruta}reportes_archived/NMAP-resumen.txt >> INTERNO/consolidado/reportes_archived/NMAP-resumen2.txt")
            os.system(f"cat {resultado_ruta}reportes_archived/reporte-OS.csv >> INTERNO/consolidado/reportes_archived/reporte-OS2.csv")
            
    os.system(f"sort INTERNO/consolidado/reportes_archived/reporte-OS2.csv | uniq > INTERNO/consolidado/reportes_archived/reporte-OS.csv")
    os.system(f"sort INTERNO/consolidado/.datos_archived/total-host-vivos2.txt | uniq > INTERNO/consolidado/.datos_archived/total-host-vivos.txt")
    os.system(f"sort INTERNO/consolidado/reportes_archived/NMAP-resumen2.txt | uniq > INTERNO/consolidado/reportes_archived/NMAP-resumen.txt")
    os.system(f"find INTERNO/consolidado/reportes_archived -size  0 -print0 |xargs -0 rm 2>/dev/null")
    mainDB_INTERNO.append('INTERNO/consolidado/.resultados.db')
    print("\n")
    # juntar las bases de datos internas a la DB principal
    merge_databases(mainDB_INTERNO[0],resultados_INTERNO_list)
else:
    mainDB_INTERNO = resultados_INTERNO_list # si solo hay un escaneo

todos_resultados = mainDB_INTERNO + resultados_EXTERNO_list
print (f'todos_resultados {todos_resultados}')
indexDB = 0
todos_resultados_len = len(todos_resultados)
#print (f'indexDB {indexDB} todos_resultados_len {todos_resultados_len}')

for resultados_db in todos_resultados:    

    print ("\n")
    indexDB = indexDB + 1
    host_analizados = 0
    #print(f"resultados_db {resultados_db}")
    ruta = resultados_db
    ruta = ruta.replace(".resultados.db", "")
    print(f'ruta {ruta}')
    vectorInforme = resultados_db.split('/')[0]   
    #time.sleep(15)
    # si no existe informe
    if not (os.path.isfile(f'{empresa}-{vectorInforme}.xlsx')):
        # crear hoja de calculo
        print (f"Generando informe {vectorInforme} para {empresa}")
        globals()['wb' + empresa + "-" +vectorInforme] = openpyxl.Workbook()
    

    if "EXTERNO" in ruta:
        dominio = ruta
        print (f'dominio {dominio}')
        dominio=dominio.replace("EXTERNO/", "")
        dominio=dominio.replace("/", "")
        resultados_EXTERNO_count = resultados_EXTERNO_count +1        
        os.system(f"mkdir EXTERNO/{dominio}/.datos_archived 2>/dev/null;cp EXTERNO/{dominio}/.datos/* EXTERNO/{dominio}/.datos_archived 2>/dev/null")
        os.system(f"mkdir EXTERNO/{dominio}/reportes_archived 2>/dev/null;cp EXTERNO/{dominio}/reportes/* EXTERNO/{dominio}/reportes_archived 2>/dev/null")
    else:
        segmento = ruta
        segmento = segmento.replace("INTERNO/", "")
        segmento = segmento.replace("/", "")
        
        
        resultados_INTERNO_count = resultados_INTERNO_count +1
    print(f"ruta {ruta}")
    print(f"dominio {dominio}")
    


    conn = sqlite3.connect(resultados_db)
    c = conn.cursor()

    #vul_externas = c.execute('select COUNT (DISTINCT TIPO) from VULNERABILIDADES').fetchone()[0]
    stream = os.popen('wc -l ' + ruta + '.datos_archived/total-host-vivos.txt | cut -d " " -f1')
    host_analizados = int(stream.read())
    total_host_analizados = total_host_analizados + host_analizados
    #print(f"total_host_analizados {total_host_analizados}")

    # Host unicos con alguna vulnerabilidad
    host_uniq_vulnerabilidades = c.execute('SELECT COUNT (DISTINCT IP) FROM VULNERABILIDADES').fetchone()[0]
    #print(f"host_uniq_vulnerabilidades {host_uniq_vulnerabilidades}")

    #mmmm
    #host_vulnerabilidades = c.execute('SELECT COUNT (IP) FROM VULNERABILIDADES;').fetchone()[0]
    #print(f"host_vulnerabilidades {host_vulnerabilidades}")

    ########################### Vulnerabilidades por criticidad #################

    # Vulnerabilidades criticas unicas
    vul_criticas = c.execute("select COUNT (DISTINCT TIPO) from VULNERABILIDADES where tipo ='passTomcat'  or tipo ='webshell' or tipo ='backdoorFabrica' or tipo ='ransomware' or tipo ='JoomlaJCKeditor'  or tipo ='HTTPsys' or tipo ='RedInterna' ").fetchone()[0]
    total_vul_criticas = total_vul_criticas + vul_criticas
    if "INTERNO" in vectorInforme: total_vul_criticas_INTERNO = total_vul_criticas_INTERNO + vul_criticas
    if "EXTERNO" in vectorInforme: total_vul_criticas_EXTERNO = total_vul_criticas_EXTERNO + vul_criticas
    #print(f"total_vul_criticas {total_vul_criticas}")


    # Vulnerabilidades altas unicas
    vul_altos = c.execute("select COUNT (DISTINCT TIPO) from VULNERABILIDADES where tipo ='ms08067' or tipo ='validacionInsuficiente' or tipo ='ms17010'  or tipo ='mailPass' or tipo ='defaultPassword' or tipo ='compartidoNFS' or tipo ='logeoRemoto' or tipo ='heartbleed'  or tipo ='passwordMikroTik' or tipo ='VNCnopass'   or tipo ='perdidaAutenticacion'  or tipo ='slowloris' or tipo ='wordpressPass' or tipo ='conficker' or tipo ='anonymousIPMI' or tipo ='noSQLDatabases' or tipo ='winboxVuln' or tipo ='rmiVuln' or tipo ='SSHBypass' or tipo ='intelVuln' or tipo ='smbrelayShell' or tipo ='backupWeb' or tipo ='apacheStruts'  or tipo ='IISwebdavVulnerable' or tipo ='shellshock' or tipo ='ciscoASAVuln' or tipo ='SambaCry' or tipo ='misfortune' or tipo ='jbossVuln'   or tipo ='poisoning' or tipo ='cipherZeroIPMI'  or tipo ='owaVul' or tipo ='hashRoto'  or tipo ='github'  or tipo ='Kerberoasting' or tipo ='passwordSFI' or tipo ='restablecimientoPassInseguro' or tipo ='sqli'  or tipo ='openrelayEXTERNO' or tipo ='rdpPass' or tipo ='ddos' or tipo ='passwordAdivinadoServ' or tipo ='passwordAdminWeb' or tipo ='passwordPhpMyadmin' or tipo ='googleRedirect' or tipo ='passwordDefecto' or tipo ='BlueKeep' or tipo ='chamilo-CVE~2023~34960' or tipo ='SMBGhost-Bleed' or tipo ='ms09050' or tipo ='malwareWeb' or tipo ='phishing' or tipo ='zimbraXXE' or tipo ='zerologon'").fetchone()[0]
    total_vul_altas = total_vul_altas + vul_altos
    if "INTERNO" in vectorInforme: total_vul_altas_INTERNO = total_vul_altas_INTERNO + vul_altos
    if "EXTERNO" in vectorInforme: total_vul_altas_EXTERNO = total_vul_altas_EXTERNO + vul_altos
    #print(f"total_vul_altas {total_vul_altas}")

    # Vulnerabilidades medias unicas
    vul_medias = c.execute("select COUNT (DISTINCT TIPO) from VULNERABILIDADES where tipo ='VPNhandshake' or tipo ='openstreaming'  or tipo ='vulnDahua' or tipo ='directorioLDAP' OR tipo LIKE 'nullsession-%' or tipo ='transferenciaDNS' or tipo ='listadoDirectorio' or tipo ='enumeracionUsuarios'  or tipo ='anonymous' or tipo ='ACL' or tipo ='ms12020'  or tipo ='exposicionUsuarios'  or tipo ='upnpAbierto'  or tipo ='wordpressPingbacks'  or tipo ='scribd' or tipo ='ftpAnonymous' or tipo ='exposicionDatosPersonales'  or tipo ='passwordAdivinadoWin' or tipo ='DoSWeb' or tipo ='openrelayINTERNO'  or tipo ='xml-rpc-habilitado' or tipo ='VNCbypass' or tipo ='archivosPeligrosos' or tipo ='webdavVulnerable' or tipo ='openWebservice' or tipo ='printNightmare' or tipo ='malwareEjecutable' or tipo like '%compartidoSMB%' or tipo ='gppPassword' or tipo ='escalamientoPrivilegios'  or tipo ='ghostcat' or tipo ='spoof' or tipo ='passwordHost' or tipo ='ListadoDirectorioArchivosSensibles' or tipo ='exfiltracionInformacion' or tipo ='fugaInformacion' or tipo ='LFI' or tipo ='hikvisionBackdoor' or tipo ='CMSDesactualizado' or tipo ='CS-10' or tipo ='auexterna-databases'").fetchone()[0]
    total_vul_medias = total_vul_medias + vul_medias
    if "INTERNO" in vectorInforme: total_vul_medias_INTERNO = total_vul_medias_INTERNO + vul_medias
    if "EXTERNO" in vectorInforme: total_vul_medias_EXTERNO = total_vul_medias_EXTERNO + vul_medias
    #print(f"total_vul_medias {total_vul_medias}")

    # Vulnerabilidades bajas unicas
    vul_bajas = c.execute("select COUNT (DISTINCT TIPO) from VULNERABILIDADES where tipo ='passwordDahua'  or tipo ='snmpCommunity'  or tipo ='listadoDirectorio' or tipo ='enumeracionUsuarios' or tipo ='googlehacking' or tipo ='anonymous' or tipo ='erroresWeb' or tipo ='ACL'  or tipo ='CVE15473'  or tipo ='MACflooding'  or tipo ='wpUsers' or tipo ='vulTLS' or tipo ='softwareObsoleto' or tipo ='nmapHTTPvuln' or tipo ='vlanHop' or tipo ='headerSegurosFaltantes' or tipo ='smbrelayShare' or tipo ='xss' or tipo ='fortinet-CVE~2023~27997' or tipo ='captcha' or tipo ='FPD' or tipo ='IPinterna' or tipo ='SourceMap' or tipo ='domainTakeOver' or tipo ='CS-50' or tipo ='confTLS' or tipo ='registroHabilitado' or tipo ='noTLS' or tipo ='passwordZKSoftware' or tipo ='passwordBD' or tipo ='hikvisio-nuclei'  or tipo ='laptopControlesInsuficientes' ").fetchone()[0] 
    total_vul_bajas = total_vul_bajas + vul_bajas
    if "INTERNO" in vectorInforme: total_vul_bajas_INTERNO = total_vul_bajas_INTERNO + vul_bajas
    if "EXTERNO" in vectorInforme: total_vul_bajas_EXTERNO = total_vul_bajas_EXTERNO + vul_bajas
    #print(f"total_vul_bajas {total_vul_bajas}")

    # Vulnerabilidades informativas unicas
    vul_info = c.execute("select COUNT (DISTINCT TIPO) from VULNERABILIDADES where tipo ='archivosDefecto'  or tipo ='divulgacionInformacion'  or tipo ='vrfyHabilitado' or tipo ='contenidoNoRelacionado' or tipo ='debugHabilitado' or tipo ='stp'  or tipo ='openresolver' or tipo ='enumeracionUsuariosSSH' or tipo ='802-1x' or tipo ='googleCachedContent' or tipo ='contenidoPrueba' or tipo ='command' or tipo ='KerberNOroasting' or tipo ='AtaqueWifi' or tipo ='ListadoDirectorios' ").fetchone()[0]
    total_vul_info = total_vul_info + vul_info
    if "INTERNO" in vectorInforme: total_vul_info_INTERNO = total_vul_info_INTERNO + vul_info
    if "EXTERNO" in vectorInforme: total_vul_info_EXTERNO = total_vul_info_EXTERNO + vul_info
    #print(f"total_vul_info {total_vul_info}")


    ###################################################


    ###### VULNERABILIDADES POR ACTIVO ####

    # aplicacionMovil
    vuln_app = c.execute("SELECT count (distinct tipo) FROM VULNERABILIDADES WHERE TIPO='validacionInsuficiente'").fetchone()[0]
    aplicacionMovil   = aplicacionMovil   + vuln_app      

    # aplicacionWeb
    vuln_app = c.execute("SELECT count (distinct tipo) FROM VULNERABILIDADES WHERE TIPO='debugHabilitado' or TIPO='ListadoDirectorios' or tipo ='ListadoDirectorioArchivosSensibles' or TIPO='archivosDefecto' or TIPO='divulgacionInformacion' or TIPO='archivosPeligrosos' or TIPO='googlehacking' or TIPO='perdidaAutenticacion' or TIPO='erroresWeb' or TIPO='wpUsers' or TIPO='exposicionUsuarios'  or TIPO='wordpressPass'  or TIPO='webshell'  or TIPO='backupWeb' or TIPO='CMSDesactualizado' or TIPO='registroHabilitado' or TIPO='captcha' or tipo ='contenidoPrueba' or tipo ='wordpressPingbacks'  or tipo ='restablecimientoPassInseguro' or tipo ='exposicionDatosPersonales' or tipo ='sqli' or tipo ='DoSWeb' or tipo ='xml-rpc-habilitado' or tipo ='googleRedirect'  or TIPO='defaultPassword' or tipo ='passwordAdminWeb' or tipo ='malwareWeb' or tipo ='xss' or tipo ='exfiltracionInformacion' or tipo ='chamilo-CVE~2023~34960' or tipo ='LFI' or tipo ='FPD' or tipo ='SourceMap' or tipo ='domainTakeOver' or tipo ='CS-10' or tipo ='CS-50'").fetchone()[0]
    aplicacionWeb   = aplicacionWeb   + vuln_app    

    #ATM
    atm = 0

    # estaciones de trabajo
    vuln_estacion = c.execute("SELECT count (distinct tipo) FROM VULNERABILIDADES WHERE  TIPO='passwordHost' or tipo like '%compartidoSMB%' or TIPO='ms17010' or TIPO='ms08067' or TIPO='BlueKeep' or TIPO='ms12020' or TIPO='doublepulsar' or TIPO='conficker' or TIPO='VNCbypass' or TIPO='VNCnopass' or TIPO='ransomware'  or tipo ='smbrelayShell'  or tipo ='passwordAdivinadoWin' or tipo ='printNightmare' or tipo ='escalamientoPrivilegios' or tipo ='smbrelayShare' or tipo ='SMBGhost-Bleed' or tipo ='ms09050' or tipo ='laptopControlesInsuficientes'").fetchone()[0]
    estacionesTrabajo    = estacionesTrabajo   + vuln_estacion    

    #WIFI
    vuln_wifi = c.execute("SELECT count (distinct tipo) FROM VULNERABILIDADES WHERE  TIPO='AtaqueWifi'").fetchone()[0]
    wifi  = wifi   + vuln_wifi

    # servidores
    vuln_serv = c.execute("SELECT count (distinct tipo) FROM VULNERABILIDADES WHERE TIPO='compartidoNFS' OR tipo LIKE 'nullsession-%' or TIPO='shellshock' or TIPO='webdavVulnerable' or TIPO='heartbleed' or TIPO='zimbraXXE' or TIPO='slowloris' or TIPO='CVE15473' or TIPO='directorioLDAP' or TIPO='transferenciaDNS' or TIPO='vrfyHabilitado' or TIPO='openresolver' or TIPO='openrelayINTERNO' or TIPO='openrelayEXTERNO' or TIPO='spoof' or TIPO='anonymousIPMI' or TIPO='rmiVuln' or TIPO='SSHBypass' or TIPO='intelVuln' or TIPO='HTTPsys' or TIPO='apacheStruts' or TIPO='IISwebdavVulnerable' or TIPO='SambaCry' or TIPO='jbossVuln' or TIPO='passwordSFI' or TIPO='contenidoNoRelacionado' or TIPO='cipherZeroIPMI' or TIPO='zerologon' or TIPO='vulTLS' or TIPO='owaVul' or TIPO='confTLS' or tipo ='noTLS'  or tipo ='ftpAnonymous' or tipo ='Kerberoasting' or TIPO='IPinterna' or tipo ='gppPassword' or tipo ='rdpPass' or tipo ='ddos' or tipo ='passwordAdivinadoServ' or tipo ='openWebservice' or tipo ='passwordDefecto' or tipo ='malwareEjecutable' or tipo ='softwareObsoleto' or tipo ='ghostcat' or tipo ='nmapHTTPvuln' or tipo ='headerSegurosFaltantes' or tipo ='enumeracionUsuariosSSH' or tipo ='command' or tipo ='KerberNOroasting'").fetchone()[0]
    servidores   = servidores   + vuln_serv

    # base de datos
    vuln_bd = c.execute("SELECT count (distinct tipo) FROM VULNERABILIDADES WHERE TIPO='passwordBD' or TIPO='noSQLDatabases' or TIPO='JoomlaJCKeditor' or tipo ='passwordPhpMyadmin' or tipo ='auexterna-databases'").fetchone()[0]
    baseDatos   = baseDatos   + vuln_bd

    # telefonia IP
    vuln_telefonia = c.execute("SELECT count (distinct tipo) FROM VULNERABILIDADES WHERE  TIPO='svmap' ").fetchone()[0]
    telefoniaIP    = telefoniaIP   + vuln_telefonia
    
    # sistema Vigilancia
    vuln_vigilancia = c.execute("SELECT count (distinct tipo) FROM VULNERABILIDADES WHERE TIPO='vulnDahua' or TIPO='openstreaming' or TIPO='passwordDahua' or TIPO='hikvisionBackdoor' or tipo ='hikvisio-nuclei'").fetchone()[0]
    sistemaVigilancia    = sistemaVigilancia   + vuln_vigilancia
    

    # Dispositivos de red
    vuln_red = c.execute("SELECT count (distinct tipo) FROM VULNERABILIDADES WHERE TIPO='passwordMikroTik' or TIPO='winboxVuln' or TIPO='snmpCommunity' or TIPO='VPNhandshake' or TIPO='backdoorFabrica' or TIPO='ciscoASAVuln' or TIPO='misfortune' or TIPO='upnpAbierto' or TIPO='poisoning' or TIPO='stp' or TIPO='vlanHop' or tipo ='MACflooding'  or tipo ='802-1x' or tipo ='fortinet-CVE~2023~27997'  or tipo ='RedInterna'  ").fetchone()[0]
    dispositivosRed    = dispositivosRed   + vuln_red        

    # personal
    vuln_personal = c.execute("SELECT count (distinct tipo) FROM VULNERABILIDADES WHERE TIPO='mailPass' or tipo ='phishing'").fetchone()[0]
    personal = personal   + vuln_personal


    # Otros
    vuln_otros = c.execute("SELECT count (distinct tipo) FROM VULNERABILIDADES WHERE tipo ='github' or tipo ='scribd' or tipo ='googleCachedContent' or tipo ='fugaInformacion' or tipo ='passwordZKSoftware' ").fetchone()[0]
    otros = otros   + vuln_otros

    

    #### VULNERABILIDADES POR CATEGORIA ####
    # password
    vuln_pass = c.execute("SELECT count (distinct tipo) FROM VULNERABILIDADES WHERE TIPO='passwordMikroTik' or TIPO='passwordAdivinadoWin' or TIPO='passwordHost' or TIPO='defaultPassword' or TIPO='mailPass' or TIPO='passwordDahua' or TIPO='passwordBD' or TIPO='noSQLDatabases' or TIPO='passwordSFI' or tipo ='hashRoto' or tipo ='rdpPass' or tipo ='phishing' or tipo ='passwordAdivinadoServ' or tipo ='Kerberoasting' or tipo ='passwordDefecto' or tipo ='passwordAdminWeb' or tipo ='passwordPhpMyadmin' or tipo ='passwordZKSoftware'").fetchone()[0]
    passwordDebil = passwordDebil + vuln_pass
    #print(f"passwordDebil  {passwordDebil}")

    # falta de parches
    vuln_parches = c.execute("SELECT count (distinct tipo) FROM VULNERABILIDADES WHERE TIPO='winboxVuln' or TIPO='shellshock' or TIPO='ms17010' or TIPO='ms08067' or TIPO='heartbleed' or TIPO='zimbraXXE' or TIPO='BlueKeep' or TIPO='slowloris' or TIPO='CVE15473' or TIPO='ms12020' or TIPO='vulnDahua'  or TIPO='doublepulsar' or TIPO='conficker' or TIPO='SSHBypass' or TIPO='VNCbypass' or TIPO='intelVuln' or TIPO='HTTPsys' or TIPO='apacheStruts' or TIPO='backdoorFabrica' or TIPO='IISwebdavVulnerable' or TIPO='ciscoASAVuln' or TIPO='SambaCry' or TIPO='misfortune' or TIPO='jbossVuln' or TIPO='cipherZeroIPMI' or TIPO='ransomware' or TIPO='JoomlaJCKeditor' or TIPO='zerologon' or TIPO='owaVul' or TIPO='CMSDesactualizado'  or TIPO='printNightmare' or TIPO='googleRedirect'  or TIPO='softwareObsoleto' or TIPO='ghostcat' or TIPO='nmapHTTPvuln' or TIPO='enumeracionUsuariosSSH' or tipo ='chamilo-CVE~2023~34960' or tipo ='fortinet-CVE~2023~27997' or tipo ='SMBGhost-Bleed' or tipo ='ms09050' or TIPO='hikvisionBackdoor' or tipo ='hikvisio-nuclei'").fetchone()[0]
    faltaParches = faltaParches + vuln_parches
    #print(f"faltaParches  {faltaParches}")
    # Errores de configuracion
    vuln_conf = c.execute("SELECT count (distinct tipo) FROM VULNERABILIDADES WHERE TIPO='logeoRemoto' or TIPO='compartidoNFS' or tipo like '%compartidoSMB%' OR tipo LIKE 'nullsession-%' or TIPO='snmpCommunity' or TIPO='directorioLDAP' or TIPO='transferenciaDNS' or TIPO='vrfyHabilitado' or TIPO='ftpAnonymous' or TIPO='openstreaming' or TIPO='VPNhandshake' or TIPO='openresolver' or TIPO='openrelayINTERNO'  or TIPO='openrelayEXTERNO' or TIPO='spoof' or TIPO='anonymousIPMI' or TIPO='rmiVuln' or TIPO='VNCnopass' or TIPO='upnpAbierto' or TIPO ='smbrelayShell' or TIPO ='smbrelayShare' or TIPO ='poisoning' or TIPO ='registroHabilitado' or TIPO='stp' or TIPO='vlanHop' or TIPO='vulTLS' or TIPO='confTLS'   or tipo ='noTLS'  or tipo ='wordpressPingbacks' or tipo ='archivosDefecto' or tipo ='ListadoDirectorios' or tipo ='archivosPeligrosos' or tipo ='scribd'  or tipo ='MACflooding' or tipo ='gppPassword' or tipo ='ddos' or tipo ='xml-rpc-habilitado' or tipo ='wpUsers' or TIPO='webdavVulnerable' or TIPO='openWebservice' or TIPO='malwareEjecutable'  or tipo ='escalamientoPrivilegios' or tipo ='googlehacking' or tipo ='malwareWeb' or tipo ='upnpAbierto' or tipo ='802-1x' or tipo ='ListadoDirectorioArchivosSensibles' or tipo ='googleCachedContent' or tipo ='RedInterna' or tipo ='auexterna-databases' or tipo ='command' or tipo ='KerberNOroasting'  or tipo ='AtaqueWifi' or tipo ='laptopControlesInsuficientes' ").fetchone()[0]
    errorConfiguracion = errorConfiguracion + vuln_conf

    # Errores de programacion
    vuln_prog = c.execute("SELECT count (distinct tipo) FROM VULNERABILIDADES WHERE TIPO='DoSWeb' or TIPO='contenidoNoRelacionado' or TIPO='captcha' or tipo ='erroresWeb' or tipo ='debugHabilitado' or tipo ='github' or tipo ='sqli' or tipo ='exposicionDatosPersonales' or tipo ='IPinterna' or tipo ='divulgacionInformacion' or tipo ='restablecimientoPassInseguro' or tipo ='xss' or tipo ='exfiltracionInformacion' or tipo ='fugaInformacion' or tipo ='LFI' or tipo ='FPD' or tipo ='SourceMap' or tipo ='domainTakeOver' or TIPO='contenidoPrueba' or tipo ='CS-10' or tipo ='CS-50'").fetchone()[0]
    programacionInsegura = programacionInsegura + vuln_prog        
    

    ### REPORTE PHISHING #    
    # print (f'vectorInforme {vectorInforme}')
    # print (f'empresa {empresa}')
    # print (f'args.phishing {args.phishing}')
    # #time.sleep(10)
    # if (vectorInforme == "EXTERNO") and (os.path.isfile(args.phishing)):
    #     print ("generando reporte phishing")

    #     globals()['wb' + empresa + "-" +vectorInforme].create_sheet()
    #     globals()['sheet' + empresa] = globals()['wb' + empresa + "-" +vectorInforme]['Sheet']
    #     globals()['sheet' + empresa].title = f'Reporte phishing'  # Change title
    #     globals()['sheet' + empresa].column_dimensions['A'].width = 22
    #     globals()['sheet' + empresa].column_dimensions['B'].width = 22
    #     globals()['sheet' + empresa].column_dimensions['C'].width = 22

    #     phishingDF = pd.read_csv('phishing.csv', sep=',', encoding='utf-8',usecols = ['email','message','details'])

    #     #mails sent
    #     mails_send_df = phishingDF[phishingDF['message'] == 'Email Sent'] 
    #     mails_send_array = mails_send_df["email"].unique()
    #     total_mail_sent = len(mails_send_array)

    #     #mails clicked
    #     mails_clicked_df = phishingDF[phishingDF['message'] == 'Clicked Link'] 
    #     mails_clicked_uniq_df = mails_clicked_df.drop_duplicates()
    #     total_mail_clicked = len(mails_clicked_uniq_df["email"])

    #     #mails opened
    #     mails_opened_df = phishingDF[phishingDF['message'] == 'Email Opened'] 
    #     frames = [mails_opened_df, mails_clicked_df]
    #     mails_opened_all_df = pd.concat(frames)
    #     del mails_opened_all_df['message']
    #     mails_opened_uniq_df = mails_opened_all_df.drop_duplicates()
    #     mails_opened_array = mails_opened_uniq_df["email"].unique()
    #     total_mail_opened = len(mails_opened_array)

    #     #Submitted Data
    #     cred_sent_df = phishingDF[phishingDF['message'] == 'Submitted Data'] 
    #     total_creds_sent = len(cred_sent_df["email"])

    #     i = 1
    #     # Cabecera correos enviados
    #     globals()['sheet' + empresa].merge_cells('A1:C1')
    #     globals()['sheet' + empresa]['A' + str(i)] = "Correos enviados"
    #     globals()['sheet' + empresa]['A' + str(i)].border = thin_border
    #     globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #     globals()['sheet' + empresa]['A' + str(i)].font = Calibri10BoldWhite
    #     globals()['sheet' + empresa]['A' + str(i)].fill = verdeFill

    #     i = i + 1
    #     #iterar
    #     columns = ['A','B','C']
    #     j = 0
    #     for email in mails_send_array:
    #         globals()['sheet' + empresa][columns[j] + str(i)] = email                    
    #         globals()['sheet' + empresa][columns[j] + str(i)].font = Calibri10                    
    #         globals()['sheet' + empresa][columns[j] + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #         globals()['sheet' + empresa][columns[j] + str(i)].border = thin_border

    #         # si ya llenamos hasta la columna C y tenemos que empezar una nueva fila
    #         if (j == 2):
    #             j = 0
    #             i = i + 1
    #         else: 
    #             j = j + 1

        
    #     # Cabecera correos abiertos
    #     i = i + 1
    #     #    globals()['sheet' + empresa].merge_cells('A1:C1')
    #     globals()['sheet' + empresa]['A' + str(i)] = "Usuarios que abrieron el correo"
    #     globals()['sheet' + empresa]['A' + str(i)].border = thin_border
    #     globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #     globals()['sheet' + empresa]['A' + str(i)].font = Calibri10BoldWhite
    #     globals()['sheet' + empresa]['A' + str(i)].fill = verdeFill
        
    #     i = i + 1
    #     for email in mails_opened_array:                        
    #         #print (f' i {i} j {j}')
    #         globals()['sheet' + empresa]['A' + str(i)] = email                    
    #         globals()['sheet' + empresa]['A' + str(i)].font = Calibri10                    
    #         globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #         globals()['sheet' + empresa]['A' + str(i)].border = thin_border

    #         i = i +1
    
    #     ##### Cabecera links clicked
    #     i = i + 1
    #     #   globals()['sheet' + empresa].merge_cells('A1:C1')
    #     globals()['sheet' + empresa]['A' + str(i)] = "Usuarios que hicieron click"
    #     globals()['sheet' + empresa]['A' + str(i)].border = thin_border
    #     globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #     globals()['sheet' + empresa]['A' + str(i)].font = Calibri10BoldWhite
    #     globals()['sheet' + empresa]['A' + str(i)].fill = verdeFill

    #     globals()['sheet' + empresa].merge_cells('B' + str(i) + ':D' + str(i))            
    #     globals()['sheet' + empresa]['B' + str(i)] = "Información recolectada"
    #     globals()['sheet' + empresa]['B' + str(i)].border = thin_border
    #     globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #     globals()['sheet' + empresa]['B' + str(i)].font = Calibri10BoldWhite
    #     globals()['sheet' + empresa]['B' + str(i)].fill = verdeFill
    #     globals()['sheet' + empresa].row_dimensions[i].height = 25

    #     i = i + 1
    #     for index, row in mails_clicked_uniq_df.iterrows():            
    #         email = row['email']
    #         details = row['details']            
    #             #print (f' i {i} j {j}')
    #         globals()['sheet' + empresa]['A' + str(i)] = email                    
    #         globals()['sheet' + empresa]['A' + str(i)].font = Calibri10                    
    #         globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #         globals()['sheet' + empresa]['A' + str(i)].border = thin_border

    #         details_json = json.loads(details)
    #         ip = details_json['browser']['address']
    #         user_agent = details_json['browser']['user-agent']
    #         globals()['sheet' + empresa].merge_cells('B' + str(i) + ':D' + str(i))                    
    #         globals()['sheet' + empresa]['B' + str(i)] = f'IP: {ip} \nNavegador: {user_agent}'                    
    #         globals()['sheet' + empresa]['B' + str(i)].font = Calibri10                    
    #         globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #         globals()['sheet' + empresa]['B' + str(i)].border = thin_border

    #         globals()['sheet' + empresa].row_dimensions[i].height = 40                    
            
    #         i = i +1

    #     ##### Cabecera datos recoletados
    #     i = i + 1
    #     #   globals()['sheet' + empresa].merge_cells('A1:C1')
    #     globals()['sheet' + empresa]['A' + str(i)] = "Correos"
    #     globals()['sheet' + empresa]['A' + str(i)].border = thin_border
    #     globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #     globals()['sheet' + empresa]['A' + str(i)].font = Calibri10BoldWhite
    #     globals()['sheet' + empresa]['A' + str(i)].fill = verdeFill

    #     globals()['sheet' + empresa].merge_cells('B' + str(i) + ':D' + str(i))            
    #     globals()['sheet' + empresa]['B' + str(i)] = "Credenciales recolectada"
    #     globals()['sheet' + empresa]['B' + str(i)].border = thin_border
    #     globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #     globals()['sheet' + empresa]['B' + str(i)].font = Calibri10BoldWhite
    #     globals()['sheet' + empresa]['B' + str(i)].fill = verdeFill
    #     globals()['sheet' + empresa].row_dimensions[i].height = 25

    #     i = i + 1
    #     for index, row in phishingDF.iterrows():
    #         event = row['message']
    #         email = row['email']
    #         details = row['details']
    #         if event == 'Submitted Data':
    #             #print (f' i {i} j {j}')
    #             globals()['sheet' + empresa]['A' + str(i)] = email                    
    #             globals()['sheet' + empresa]['A' + str(i)].font = Calibri10                    
    #             globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #             globals()['sheet' + empresa]['A' + str(i)].border = thin_border

    #             details_json = json.loads(details)
    #             #print(json.dumps(details_json, indent=4))
    #             #time.sleep(10)

    #             username = email #details_json['payload']['login'][0]
    #             password = details_json['payload']['password'][0]
    #             globals()['sheet' + empresa].merge_cells('B' + str(i) + ':D' + str(i))                    
    #             globals()['sheet' + empresa]['B' + str(i)] = f'Usuario {username} \nPassword: {password}'                    
    #             globals()['sheet' + empresa]['B' + str(i)].font = Calibri10                    
    #             globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #             globals()['sheet' + empresa]['B' + str(i)].border = thin_border

    #             globals()['sheet' + empresa].row_dimensions[i].height = 40                    
                
    #             i = i +1
        
    #     i = i +2
    #     # Estadisticas phishing
    #     globals()['sheet' + empresa]['A' + str(i)] = "Evento"
    #     globals()['sheet' + empresa]['A' + str(i)].border = thin_border
    #     globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #     globals()['sheet' + empresa]['A' + str(i)].font = Calibri10BoldWhite
    #     globals()['sheet' + empresa]['A' + str(i)].fill = verdeFill
        
    #     globals()['sheet' + empresa]['B' + str(i)] = "Incidencias"
    #     globals()['sheet' + empresa]['B' + str(i)].border = thin_border
    #     globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #     globals()['sheet' + empresa]['B' + str(i)].font = Calibri10BoldWhite
    #     globals()['sheet' + empresa]['B' + str(i)].fill = verdeFill
    #     globals()['sheet' + empresa].row_dimensions[i].height = 25

        
    #     i = i +1
    #     globals()['sheet' + empresa]['A' + str(i)] = 'Correos enviados'
    #     globals()['sheet' + empresa]['B' + str(i)] = total_mail_sent
    #     globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #     i = i +1
    #     globals()['sheet' + empresa]['A' + str(i)] = 'Correos abiertos'
    #     globals()['sheet' + empresa]['B' + str(i)] = total_mail_opened
    #     globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #     i = i +1
    #     globals()['sheet' + empresa]['A' + str(i)] = 'Clicks de usuarios'
    #     globals()['sheet' + empresa]['B' + str(i)] = total_mail_clicked
    #     globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    #     i = i +1
    #     globals()['sheet' + empresa]['A' + str(i)] = 'Credenciales enviadas'
    #     globals()['sheet' + empresa]['B' + str(i)] = total_creds_sent
    #     globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

    #     ###grafica piiii
    #     chart = BarChart()
    #     chart.type = "col"
    #     chart.style = 10
    #     chart.shape = 4
    #     chart.title = "Estadisticas del ataque de phishing"
    #     start_data = i-3
    #     print (f'start_data {start_data}')

    #     # create data for plotting
    #     labels = Reference( globals()['sheet' + empresa], min_col=1, min_row=start_data, max_row=start_data + 3)
    #     data = Reference( globals()['sheet' + empresa], min_col=2, min_row=start_data, max_row=start_data + 3)

    #     # adding data to the Doughnut chart object
    #     chart.add_data(data, titles_from_data=False)
    #     chart.set_categories(labels)

    #     chart.dataLabels = DataLabelList()
    #     chart.dataLabels.showPercent = True
    #     chart.dataLabels.showVal = False
    #     chart.dataLabels.showLegendKey = True
    #     chart.dataLabels.showCatName = False

    #     # try to set color blue (0000FF) for the 2nd wedge (idx=1) in the series
    #     series = chart.series[0]
    #     pt = DataPoint(idx=0)
    #     pt.graphicalProperties.solidFill = "95B3D7"
    #     series.dPt.append(pt)

    #     pt = DataPoint(idx=1)
    #     pt.graphicalProperties.solidFill = "FFFF00"
    #     series.dPt.append(pt)

    #     pt = DataPoint(idx=2)
    #     pt.graphicalProperties.solidFill = "FFC000"
    #     series.dPt.append(pt)

    #     pt = DataPoint(idx=3)
    #     pt.graphicalProperties.solidFill = "FF0000"
    #     series.dPt.append(pt)


    #     #adicionar la grafica a la hoja de calculo
    #     globals()['sheet' + empresa].add_chart(chart, 'A' + str(i+2))

        


        
    ###### SUBDOMINIOS ####        
    if (os.path.isfile(ruta + "importarMaltego/subdominios.csv")):
        print (f'elaborar reporte de subdominio ')
        
        subdomainDF = pd.read_csv(ruta + "importarMaltego/subdominios.csv", sep=',', encoding='utf-8', header=None)
        # print(nmapDF.head())
        globals()['wb' + empresa + "-" +vectorInforme].create_sheet()
        globals()['sheet' + empresa] = globals()['wb' + empresa + "-" +vectorInforme]['Sheet']
        globals()['sheet' + empresa].title = f'{dominio}'  # Change title
        globals()['sheet' + empresa].column_dimensions['A'].width = 18
        globals()['sheet' + empresa].column_dimensions['B'].width = 18
        globals()['sheet' + empresa].column_dimensions['C'].width = 18
        globals()['sheet' + empresa].column_dimensions['D'].width = 18

        i = 1
        # Cabecera
        i = i + 1
        globals()['sheet' + empresa]['A' + str(i)] = "Subdominio"
        globals()['sheet' + empresa]['A' + str(i)].border = thin_border
        globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['A' + str(i)].font = Calibri10BoldWhite
        globals()['sheet' + empresa]['A' + str(i)].fill = verdeFill

        globals()['sheet' + empresa]['B' + str(i)] = "IP"
        globals()['sheet' + empresa]['B' + str(i)].border = thin_border
        globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['B' + str(i)].font = Calibri10BoldWhite
        globals()['sheet' + empresa]['B' + str(i)].fill = verdeFill

        globals()['sheet' + empresa]['C' + str(i)] = "ISP"
        globals()['sheet' + empresa]['C' + str(i)].border = thin_border
        globals()['sheet' + empresa]['C' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['C' + str(i)].font = Calibri10BoldWhite
        globals()['sheet' + empresa]['C' + str(i)].fill = verdeFill

        globals()['sheet' + empresa]['D' + str(i)] = "Ubicación"
        globals()['sheet' + empresa]['D' + str(i)].border = thin_border
        globals()['sheet' + empresa]['D' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['D' + str(i)].font = Calibri10BoldWhite
        globals()['sheet' + empresa]['D' + str(i)].fill = verdeFill

        i = i + 1
        for index, row in subdomainDF.iterrows():
            ip = row[1]
            subdominio = row[2]
            ubicacion = row[3]
            isp = str(row[4])
            # print(f"os {sistemaOperativo}")
            # print(f"os type {type(sistemaOperativo)}")

            globals()['sheet' + empresa]['A' + str(i)] = ip
            globals()['sheet' + empresa]['B' + str(i)] = subdominio
            globals()['sheet' + empresa]['C' + str(i)] = ubicacion
            globals()['sheet' + empresa]['D' + str(i)] = isp

            globals()['sheet' + empresa]['A' + str(i)].font = Calibri10
            globals()['sheet' + empresa]['B' + str(i)].font = Calibri10
            globals()['sheet' + empresa]['C' + str(i)].font = Calibri10
            globals()['sheet' + empresa]['D' + str(i)].font = Calibri10

            globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
            globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
            globals()['sheet' + empresa]['C' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
            globals()['sheet' + empresa]['D' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

            globals()['sheet' + empresa]['A' + str(i)].border = thin_border
            globals()['sheet' + empresa]['B' + str(i)].border = thin_border
            globals()['sheet' + empresa]['C' + str(i)].border = thin_border
            globals()['sheet' + empresa]['D' + str(i)].border = thin_border
            i = i + 1
        
        # Agregar Whois
        logPath = f'{ruta}logs/enumeracion/{dominio}_dns_whois.txt'
        log2Image(logPath,'whois.png')
        imgVul = openpyxl.drawing.image.Image('whois.png')            
        globals()['sheet' + empresa].add_image(imgVul, "F1")

        # Agregar dnsenum
        logPath = f'{ruta}logs/enumeracion/dnsenum.txt'
        log2Image(logPath,'dnsenum.png')
        imgVul = openpyxl.drawing.image.Image('dnsenum.png')            
        globals()['sheet' + empresa].add_image(imgVul, "O1")



    ###### PUERTOS ABIERTOS ####
    #/root/AUDITORIAS2021/GAINZA/EXTERNO/gainza.com.bo/logs/enumeracion/
    nmapDF = pd.read_csv(ruta + "reportes_archived/NMAP-resumen.txt", sep='\t', encoding='utf-8', header=None)
    scanned_ips = []
    #print(nmapDF.head())
    globals()['wb' + empresa + "-" +vectorInforme].create_sheet()
    globals()['sheet' + empresa] = globals()['wb' + empresa + "-" +vectorInforme]['Sheet']
    globals()['sheet' + empresa].title = f'Puertos abiertos ({segmento})'  # Change title
    globals()['sheet' + empresa].column_dimensions['A'].width = 15
    globals()['sheet' + empresa].column_dimensions['B'].width = 25
    globals()['sheet' + empresa].column_dimensions['C'].width = 15

    i = 1
    # Cabecera
    i = i + 1
    globals()['sheet' + empresa].merge_cells('B' + str(i) + ':C' + str(i))
    globals()['sheet' + empresa]['A' + str(i)] = "IP"
    globals()['sheet' + empresa]['A' + str(i)].border = thin_border
    globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    globals()['sheet' + empresa]['A' + str(i)].font = Calibri10BoldWhite
    globals()['sheet' + empresa]['A' + str(i)].fill = verdeFill

    globals()['sheet' + empresa]['B' + str(i)] = "Puertos abiertos"
    globals()['sheet' + empresa]['B' + str(i)].border = thin_border
    globals()['sheet' + empresa]['C' + str(i)].border = thin_border
    globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    globals()['sheet' + empresa]['B' + str(i)].font = Calibri10BoldWhite
    globals()['sheet' + empresa]['B' + str(i)].fill = verdeFill

    i = i + 1
    for index, row in nmapDF.iterrows():
        ip = row[0]
        tcp = row[1]
        udp = row[2]

        if ip not in scanned_ips :
            globals()['sheet' + empresa]['A' + str(i)] = ip
            globals()['sheet' + empresa]['B' + str(i)] = tcp
            globals()['sheet' + empresa]['C' + str(i)] = udp

            globals()['sheet' + empresa]['A' + str(i)].font = Calibri10
            globals()['sheet' + empresa]['B' + str(i)].font = Calibri10
            globals()['sheet' + empresa]['C' + str(i)].font = Calibri10

            globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
            globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
            globals()['sheet' + empresa]['C' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

            globals()['sheet' + empresa]['A' + str(i)].border = thin_border
            globals()['sheet' + empresa]['B' + str(i)].border = thin_border
            globals()['sheet' + empresa]['C' + str(i)].border = thin_border
            i = i + 1
        scanned_ips.append(ip)

    ###### SISTEMAS OPERATIVOS ####
    if (os.path.isfile(ruta + "reportes_archived/reporte-OS.csv")):
        osDF = pd.read_csv(ruta + "reportes_archived/reporte-OS.csv", sep='|', encoding='utf-8',header=None)
        print(osDF)
        # print(nmapDF.head())
        globals()['wb' + empresa + "-" +vectorInforme].create_sheet()
        globals()['sheet' + empresa] = globals()['wb' + empresa + "-" +vectorInforme]['Sheet']
        globals()['sheet' + empresa].title = 'Sistemas operativos'  # Change title
        globals()['sheet' + empresa].column_dimensions['A'].width = 15
        globals()['sheet' + empresa].column_dimensions['B'].width = 40
        globals()['sheet' + empresa].column_dimensions['C'].width = 15
        globals()['sheet' + empresa].column_dimensions['D'].width = 30

        i = 1
        # Cabecera
        i = i + 1
        globals()['sheet' + empresa]['A' + str(i)] = "IP"
        globals()['sheet' + empresa]['A' + str(i)].border = thin_border
        globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['A' + str(i)].font = Calibri10BoldWhite
        globals()['sheet' + empresa]['A' + str(i)].fill = verdeFill
        
        globals()['sheet' + empresa]['B' + str(i)] = "S.O."
        globals()['sheet' + empresa]['B' + str(i)].border = thin_border
        globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['B' + str(i)].font = Calibri10BoldWhite
        globals()['sheet' + empresa]['B' + str(i)].fill = verdeFill

        globals()['sheet' + empresa]['C' + str(i)] = "Dominio"
        globals()['sheet' + empresa]['C' + str(i)].border = thin_border
        globals()['sheet' + empresa]['C' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['C' + str(i)].font = Calibri10BoldWhite
        globals()['sheet' + empresa]['C' + str(i)].fill = verdeFill

        globals()['sheet' + empresa]['D' + str(i)] = "Nombre"
        globals()['sheet' + empresa]['D' + str(i)].border = thin_border
        globals()['sheet' + empresa]['D' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['D' + str(i)].font = Calibri10BoldWhite
        globals()['sheet' + empresa]['D' + str(i)].fill = verdeFill

        i = i + 1
        for index, row in osDF.iterrows():
            ip = row[0]
            sistemaOperativo = row[1]
            #dominioRed = row[2]
            #nombre = str(row[4])
            print(f"os ({sistemaOperativo})")
            # print(f"os type {type(sistemaOperativo)}")
            #if math.isnan(sistemaOperativo):
            if type(sistemaOperativo) != str :
                print ("OS vacio")
                continue
            print ("OS row")
            globals()['sheet' + empresa]['A' + str(i)] = ip
            globals()['sheet' + empresa]['B' + str(i)] = sistemaOperativo
            #globals()['sheet' + empresa]['C' + str(i)] = dominioRed
            #globals()['sheet' + empresa]['D' + str(i)] = sistemaOperativo

            globals()['sheet' + empresa]['A' + str(i)].font = Calibri10
            globals()['sheet' + empresa]['B' + str(i)].font = Calibri10
            #globals()['sheet' + empresa]['C' + str(i)].font = Calibri10
            #globals()['sheet' + empresa]['D' + str(i)].font = Calibri10

            globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
            globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
            #globals()['sheet' + empresa]['C' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
            #globals()['sheet' + empresa]['D' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

            globals()['sheet' + empresa]['A' + str(i)].border = thin_border
            globals()['sheet' + empresa]['B' + str(i)].border = thin_border
            #globals()['sheet' + empresa]['C' + str(i)].border = thin_border
            #globals()['sheet' + empresa]['D' + str(i)].border = thin_border
            i = i + 1

    

    ##### PRUEBAS REALIZADAS ####
    #print("Generando reporte de pruebas realizadas")
    globals()['wb' + empresa + "-" +vectorInforme].create_sheet()
    globals()['sheet' + empresa] = globals()['wb' + empresa + "-" +vectorInforme]['Sheet']
    globals()['sheet' + empresa].title = 'pruebas'  # Change title.
    globals()['sheet' + empresa].column_dimensions['A'].width = 15
    globals()['sheet' + empresa].column_dimensions['B'].width = 10
    globals()['sheet' + empresa].column_dimensions['C'].width = 15
    globals()['sheet' + empresa].column_dimensions['D'].width = 48

    # Num pruebas de vulnerabilidades
    stream = os.popen('ls ' + ruta + '/logs/vulnerabilidades | cut -d "_" -f3 | sort | uniq | wc -l')
    pruebasVulnerabilidades_cant = int(stream.read())

    # Num pruebas de cracking
    stream = os.popen('ls ' + ruta + '/logs/cracking | cut -d "_" -f3 | sort | uniq | wc -l')
    pruebasPassword_cant = int(stream.read())

    totalPruebas = totalPruebas + pruebasVulnerabilidades_cant + pruebasPassword_cant
    #print (f'ls {ruta}/logs/vulnerabilidades')
    #print(f"totalPruebas {totalPruebas}")

    # pruebas de vulnerabilidades
    stream = os.popen('ls ' + ruta + '/logs/vulnerabilidades')
    pruebasVulnerabilidades = stream.read()
    pruebasVulnerabilidades_list = pruebasVulnerabilidades.split('\n')
    pruebasVulnerabilidades_list.pop()
    # print(f"pruebasVulnerabilidades {pruebasVulnerabilidades_list}")

    # pruebas de Cracking
    stream = os.popen('ls ' + ruta + '/logs/cracking/| grep --color=never ".txt"')
    pruebasCracking = stream.read()
    pruebasCracking_list = pruebasCracking.split('\n')
    pruebasCracking_list.pop()
    # print(f"pruebasCracking {pruebasCracking_list}")

    # pruebas de Enumeracion
    stream = os.popen('ls ' + ruta + '/logs/enumeracion/ | grep --color=never web | egrep -v "Data|canary|wget"')
    pruebasEnumeracion = stream.read()
    pruebasEnumeracion_list = pruebasEnumeracion.split('\n')
    pruebasEnumeracion_list.pop()
    # print(f"pruebasEnumeracion {pruebasEnumeracion_list}")

    todasPruebas = pruebasVulnerabilidades_list + pruebasCracking_list + pruebasEnumeracion_list
    #print(f"todasPruebas {todasPruebas}")

    stream = os.popen('cat ' + ruta + 'passwords.txt | wc -l')
    totalPasswords = int(stream.read())  # totalPasswords probados
    # print(f"totalPasswords {totalPasswords}")

    webApp_tests = []
    servers_tests = []
    bd_tests = []
    workstation_tests = []
    video_tests = []
    network_tests = []
    passwords_tests = []
    testHacking = []

    for prueba in todasPruebas:
        print(f"prueba {prueba}")
        prueba_list = prueba.split('_')
        ip = prueba_list[0]
        port = prueba_list[1]
        vuln = prueba_list[2]
        vuln = vuln.replace(".txt", "")
        vuln = vuln.replace(".html", "")        
        cod_desc = buscarDescripcion(vuln)
        CVE = cod_desc.split(";")[0]
        desc = cod_desc.split(";")[1]

        desc = desc.replace("DOMINIO", dominio)
        # print(f"vuln {vuln} ip {ip} port {port} CVE {CVE} desc {desc}")

        # Activos de informacion - aplicaciones web
        if vuln in ['debugHabilitado', 'ListadoDirectorios', 'divulgacionInformacion', 'archivosDefecto',
                    'archivosPeligrosos',  'wpUsers', 'perdidaAutenticacion', 'exposicionUsuarios',
                    'wordpressPass', 'backupWeb', 'CMSDesactualizado', 'webshell','xss',
                    'registroHabilitado','googleRedirect','contenidoPrueba', 'JoomlaJCKeditor','github', 'scribd','sqli','DoSWeb','mailPass','xml-rpc-habilitado','captcha','defaultPassword','passwordAdminWeb','malwareWeb','exfiltracionInformacion','chamilo-CVE~2023~34960','googleCachedContent','fugaInformacion','LFI','FPD','SourceMap','domainTakeOver']:
            webApp_tests.append({'ip': ip, 'port': port, 'CVE': CVE, 'desc': desc})
        if 'googlehacking' in vuln:
            webApp_tests.append({'ip': ip, 'port': "N/A", 'CVE': CVE, 'desc': desc})

        if 'password' in vuln:
            passwords_tests.append({'ip': ip, 'port': port, 'CVE': CVE, 'desc': desc})            

        # Activos de informacion - servidores
        if vuln in ['compartidoNFS', 'nullsession', 'shellshock', 'webdavVulnerable', 'heartbleed', 'zimbraXXE',
                    'slowloris', 'CVE15473', 'directorioLDAP', 'transferenciaDNS', 'vrfyHabilitado', 'openresolver',
                    'openrelayINTERNO', 'openrelayEXTERNO', 'anonymousIPMI', 'rmiVuln', 'SSHBypass', 'intelVuln', 'HTTPsys', 'apacheStruts',
                    'IISwebdavVulnerable', 'SambaCry', 'jbossVuln', 'contenidoNoRelacionado', 'spoof',
                    'cipherZeroIPMI', 'erroresWeb','IPinterna','rdpPass','ddos',
                    'zerologon', 'vulTLS', 'confTLS', 'owaVul', 'noTLS', 'ftpAnonymous','passwordHost', 'Kerberoasting','passwordSFI','restablecimientoPassInseguro','exposicionDatosPersonales', 'gppPassword','passwordBD','passwordAdivinadoServ','openWebservice','passwordDefecto','passwordPhpMyadmin','malwareEjecutable','softwareObsoleto','ghostcat','nmapHTTPvuln','headerSegurosFaltantes','KerberNOroasting']:
            servers_tests.append({'ip': ip, 'port': port, 'CVE': CVE, 'desc': desc})

        # Activos de informacion - baseDatos
        if vuln in ['noSQLDatabases']:
            bd_tests.append({'ip': ip, 'port': port, 'CVE': CVE, 'desc': desc})

        # Activos de informacion - estacionesTrabajo
        if vuln in ['compartidoSMB', 'ms17010', 'ms08067', 'BlueKeep', 'ms12020', 'doublepulsar', 'conficker',
                    'VNCbypass', 'VNCnopass', 'ransomware', 'smbrelayShell','smbrelayShare', 'hashRoto','passwordAdivinadoWin', 'printNightmare','escalamientoPrivilegios','SMBGhost-Bleed','ms09050']:
            workstation_tests.append({'ip': ip, 'port': port, 'CVE': CVE, 'desc': desc})

        # Activos de informacion - sistemaVigilancia
        if vuln in ['vulnDahua', 'openstreaming', 'passwordDahua','hikvisionBackdoor']:
            video_tests.append({'ip': ip, 'port': port, 'CVE': CVE, 'desc': desc})

        # Activos de informacion - dispositivosRed
        if vuln in ['winboxVuln', 'snmpCommunity', 'VPNhandshake', 'backdoorFabrica', 'ciscoASAVuln', 'misfortune',
                    'upnpAbierto', 'poisoning', 'stp', 'vlanHop','802-1x','MACflooding','fortinet-CVE~2023~27997','RedInterna']:
            network_tests.append({'ip': ip, 'port': port, 'CVE': CVE, 'desc': desc})



    testHacking.append({'title': "Pruebas a aplicaciones web", 'tests': webApp_tests})
    testHacking.append({'title': "Pruebas a servidores (web, SMB, correo, etc)", 'tests': servers_tests})
    testHacking.append({'title': "Pruebas a base de datos", 'tests': bd_tests})
    testHacking.append({'title': "Pruebas a estaciones de trabajo", 'tests': workstation_tests})
    testHacking.append({'title': "Pruebas a sistemas de vigilancia ", 'tests': video_tests})
    testHacking.append({'title': "Pruebas a dispositivos de red", 'tests': network_tests})
    testHacking.append(
        {'title': f"Pruebas de password a servicios y dispositivos (Passwords probados {totalPasswords})",
            'tests': passwords_tests})    

    i = 1
    for details in testHacking:        
        testTitle = details['title']
        test_category = details['tests']

        # Cabecera
        i = i + 1
        globals()['sheet' + empresa].merge_cells('A' + str(i) + ':D' + str(i))
        globals()['sheet' + empresa]['A' + str(i)] = testTitle
        globals()['sheet' + empresa]['A' + str(i)].border = thin_border
        globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['A' + str(i)].font = Calibri10BoldWhite
        globals()['sheet' + empresa]['A' + str(i)].fill = verdeFill
        i = i + 1

        globals()['sheet' + empresa]['A' + str(i)] = "IP"
        globals()['sheet' + empresa]['B' + str(i)] = "Puerto"
        globals()['sheet' + empresa]['C' + str(i)] = "Código CVE"
        globals()['sheet' + empresa]['D' + str(i)] = "Prueba realizada"

        globals()['sheet' + empresa]['A' + str(i)].border = thin_border
        globals()['sheet' + empresa]['B' + str(i)].border = thin_border
        globals()['sheet' + empresa]['C' + str(i)].border = thin_border
        globals()['sheet' + empresa]['D' + str(i)].border = thin_border

        globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['C' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['D' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

        globals()['sheet' + empresa]['A' + str(i)].font = Calibri10BoldWhite
        globals()['sheet' + empresa]['B' + str(i)].font = Calibri10BoldWhite
        globals()['sheet' + empresa]['C' + str(i)].font = Calibri10BoldWhite
        globals()['sheet' + empresa]['D' + str(i)].font = Calibri10BoldWhite

        globals()['sheet' + empresa]['A' + str(i)].fill = verdeFill
        globals()['sheet' + empresa]['B' + str(i)].fill = verdeFill
        globals()['sheet' + empresa]['C' + str(i)].fill = verdeFill
        globals()['sheet' + empresa]['D' + str(i)].fill = verdeFill

        i = i + 1
        for test in test_category:
            ip = test['ip']
            port = test['port']
            CVE = test['CVE']
            desc = test['desc']

            globals()['sheet' + empresa]['A' + str(i)] = ip
            globals()['sheet' + empresa]['B' + str(i)] = port
            globals()['sheet' + empresa]['C' + str(i)] = CVE
            globals()['sheet' + empresa]['D' + str(i)] = desc

            globals()['sheet' + empresa]['A' + str(i)].font = Calibri10
            globals()['sheet' + empresa]['B' + str(i)].font = Calibri10
            globals()['sheet' + empresa]['C' + str(i)].font = Calibri10
            globals()['sheet' + empresa]['D' + str(i)].font = Calibri10

            globals()['sheet' + empresa]['A' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
            globals()['sheet' + empresa]['B' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
            globals()['sheet' + empresa]['C' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
            globals()['sheet' + empresa]['D' + str(i)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

            globals()['sheet' + empresa]['A' + str(i)].border = thin_border
            globals()['sheet' + empresa]['B' + str(i)].border = thin_border
            globals()['sheet' + empresa]['C' + str(i)].border = thin_border
            globals()['sheet' + empresa]['D' + str(i)].border = thin_border

            i = i + 1

    ######################  CREAR FICHAS DE VULNERABILIDADES ############
    # si no existe informe EXTERNO/INTERNO seteamos a 1
    if not (os.path.isfile(f'{empresa}-{vectorInforme}.xlsx')):
        numeroVulnerabilidad = 1

    
    for i in range(total_vulnerabilidades_xml):        
        vulnerabilidadImagen = 0                
        codVul = root[i].find("codVul").text        
        print (f'codVul : {codVul}')  
        conclusion = root[i].find("conclusion").text
        recomendacionMatriz = root[i].find("recomendacionMatriz").text        
        nombre = root[i].find("nombre").text
        #print("nombre " + nombre)
        CVE = root[i].find("CVE").text        
        descripcion = root[i].find("descripcion").text
        descripcion = descripcion.replace("SALTOLINEA", '\n')  
        PruebaConcepto = root[i].find("PruebaConcepto").text
        activo = root[i].find("activo").text
        PruebaConcepto = PruebaConcepto.replace("DOMINIOENTIDAD", dominio)
        vector = root[i].find("vector_"+vectorInforme).text
        score = root[i].find("score_"+vectorInforme).text
        riesgoInforme = root[i].find("riesgoInforme_"+vectorInforme).text
        agente_amenaza = root[i].find("agente_amenaza").text
        impacto_tecnico = root[i].find("impacto_tecnico").text
        impacto_negocio = root[i].find("impacto_negocio").text
        referenciaweb = root[i].find("referenciaweb").text
        referenciaweb = referenciaweb.replace("SALTOLINEA", '\n')  
        recomendacion = root[i].find("recomendacion").text

        recomendacion = recomendacion.replace("DOMINIOENTIDAD", dominio)
        recomendacion = recomendacion.replace("AMPERSAND", '&')        
        ATT_Tactic = root[i].find("ATT_Tactic").text  
        condicionesPrevias = 'Ninguna'
        try:
            condicionesPrevias = root[i].find("condicionesPrevias").text          
        except:
            error = ""
            #print("No conclu")
        
        #sql = "SELECT distinct *  FROM VULNERABILIDADES WHERE TIPO =\"" + codVul + "\";"
        sql = "SELECT distinct *  FROM VULNERABILIDADES WHERE TIPO like \"%" + codVul + "%\";"
        resSQL = c.execute(sql)
        filas = 1
        hosts = ""

        # ACCOUNT FOUND: [postgres] Host: 192.168.2.222 User: postgres Password:  [SUCCESS]
        # [MongoDB] $respuesta
        # [Redis] $respuesta"        
                #mostra IP y detalle de la vulnerabilidad
        if codVul in ['archivosPeligrosos', 'archivosDefecto', 'perdidaAutenticacion', 'webshell', 'backupWeb',
                    'ciscoASAVuln',
                    'poisoning', 'captcha', 'noTLS',  'CMSDesactualizado', 'wpUsers','erroresWeb','debugHabilitado','divulgacionInformacion','IPinterna','sqli','gppPassword','nullsession-localUsers','nullsession-domainUsers','mailPass','phishing','xml-rpc-habilitado','captcha','contenidoNoRelacionado','googleRedirect','defaultPassword','compartidoNFS','zimbraXXE','webdavVulnerable','Kerberoasting','openWebservice','softwareObsoleto','ghostcat','googlehacking','passwordSFI','malwareWeb','headerSegurosFaltantes','xss','fortinet-CVE~2023~27997','googleCachedContent','LFI','domainTakeOver','registroHabilitado','contenidoPrueba','shellshock','validacionInsuficiente','KerberNOroasting','passwordZKSoftware']:
            #print (f'codVul {codVul}')
            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]

                
                
                #Convertir salida de herramienta a imagen
                if codVul not in vuln_no_log: # si tiene log
                    if codVul == 'KerberNOroasting':
                        logPath = f'{ruta}logs/vulnerabilidades/{ip}_{port}_Kerberoasting.txt'
                    else:
                        logPath = f'{ruta}logs/vulnerabilidades/{ip}_{port}_{vulType}.txt'
                    print(logPath)
                    log2Image(logPath,codVul+'.png')  

               

                vuln_detalles = row[3]
                vuln_detalles = vuln_detalles.replace("Mensaje de error", "")
                vuln_detalles = vuln_detalles.replace("Posible Backdoor", "")
                vuln_detalles = vuln_detalles.replace("200", "")
                vuln_detalles = vuln_detalles.replace("TRACE", "")

                # https://sigec.fonadin.gob.bo:443/.git/	,
                print(f'ea {hosts}  {ip} {port} {vuln_detalles}')
                hosts = hosts + ip +":"+ port + "\n" + vuln_detalles + "\n"
                filas = filas + 1

        # Solo se muestra la IP y el puerto en el campo "host"
        if codVul in ['ms17010', 'ms08067', 'passwordDahua', 'heartbleed', 'directorioLDAP', 'spoof', 'ftpAnonymous',
                    'anonymousIPMI', 'openstreaming', 'VPNhandshake',
                    'BlueKeep', 'slowloris', 'openresolver', 'CVE15473', 'ms12020', 'doublepulsar',
                    'conficker','802-1x',
                    'rmiVuln', 'SSHBypass', 'VNCnopass', 'intelVuln',
                    'HTTPsys', 'apacheStruts', 'IISwebdavVulnerable', 'SambaCry', 'misfortune', 'jbossVuln',
                    'upnpAbierto', 'ransomware', 'cipherZeroIPMI', 'stp',
                    'vlanHop', 'JoomlaJCKeditor', 'zerologon', 'wordpressPingbacks','printNightmare','vulnDahua','restablecimientoPassInseguro','exposicionDatosPersonales','MACflooding','DoSWeb','openrelayINTERNO', 'openrelayEXTERNO','ddos','VNCbypass','googleRedirect','malwareEjecutable','escalamientoPrivilegios','nmapHTTPvuln','enumeracionUsuariosSSH','smbrelayShare','exfiltracionInformacion','chamilo-CVE~2023~34960','SMBGhost-Bleed','ms09050','fugaInformacion','FPD','hikvisionBackdoor','SourceMap','RedInterna','CS-10','CS-50','auexterna-databases','command','hikvisio-nuclei','AtaqueWifi','laptopControlesInsuficientes']:

            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]
                
                if codVul not in vuln_no_log: # si tiene log
                    #Convertir salida de herramienta a imagen
                    if "printNightmare" in codVul:
                        logPath = f'{ruta}.vulnerabilidades2/{ip}_{port}_{vulType}.txt' #generar imagen del log corto
                    else:
                        logPath = f'{ruta}logs/vulnerabilidades/{ip}_{port}_{vulType}.txt'
                    print(f'codVul {codVul} logPath {logPath}')
                    log2Image(logPath,codVul+'.png') 
                
                hosts = hosts + ip + "\n"
                filas = filas + 1
                # hosts = hosts + row[0] + "\n<br>"
                # if (( $filas % 5 == 0 ) && ($filas >0))
            # { $hosts = $hosts.$row[0]."\n<br>"; #$hosts = $hosts."<td>".$row[0]."</td></tr><tr>";}
            # else
            # { $hosts = $hosts.$row[0]."&nbsp;&nbsp;&nbsp;"; #$hosts = $hosts."<td>".$row[0]."</td>";}
            # $filas++;                            

        if codVul in ['passwordBD', 'rdpPass','defaultPassword']:            
            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]
                
                #Convertir salida de herramienta a imagen
                if codVul not in vuln_no_log: # si tiene log
                    logPath = f'{ruta}logs/cracking/{ip}_{port}_{vulType}.txt'
                    log2Image(logPath,codVul+'.png')                   

                vuln_detalles = row[3]
                vuln_detalles = vuln_detalles.replace("User", "Usuario")
                vuln_detalles = vuln_detalles.replace("ACCOUNT FOUND:", "")
                vuln_detalles = vuln_detalles.replace("[SUCCESS]", "")
                vuln_detalles = vuln_detalles.replace("Host", "IP")
                vuln_detalles = vuln_detalles.replace("Password encontrado:", "")
                hosts = hosts + vuln_detalles + "\n"
                filas = filas + 1


        if ("smbrelayShell" in codVul):
            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]
                
                #Convertir salida de herramienta a imagen
                #log2Image(logPath,codVul+'.png') 

                vuln_detalles = row[3].split(':')
                respuesta_ntlm = ':'.join([vuln_detalles[0], vuln_detalles[2]])  # 
                hosts = hosts + ip  + ": " + respuesta_ntlm + "\n"
                filas = filas + 1

        if ("ListadoDirectorios" == codVul or 'ListadoDirectorioArchivosSensibles' == codVul):
            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]
                print(f'{ip} {port}')
                
                #Convertir salida de herramienta a imagen
                #log2Image(logPath,codVul+'.png')  

                vuln_detalles = row[3]
                # $vuln_detalles =~ s/\n/<br>- /g;
                # $vuln_detalles =~ s/-  \n/<br>- /g;
                if ("index" in vuln_detalles):
                    if port in ['80', '81', '82', '83', '84', '85', '86', '8080', '8081', '8082', '8010',
                                '8800']:  hosts = hosts + f"- http://{ip}:{port} \n"
                    if port in ['443', '8443', '4443', '4433']:  hosts = hosts + f"- https://{ip}:{port} \n"
                else:
                    vuln_detalles = vuln_detalles.replace("200 ", "")
                    hosts = hosts + vuln_detalles + "\n"
                filas = filas + 1

            # $hosts = "- ".$hosts;

        if ("confTLS" in codVul):
            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]
                
                #Convertir salida de herramienta a imagen
                logPath = f'{ruta}logs/vulnerabilidades/{ip}_{port}_{vulType}.txt'
                log2Image(logPath,codVul+'.png')  

                vuln_detalles = row[3]
                vuln_detalles = vuln_detalles.replace("Configuracion incorrecta: ", "")
                vuln_detalles = vuln_detalles.replace("\n", " ,")

                hosts = hosts + f"\n-{ip}:{port} : "

                if ("TLS 1.0 habilitado" in vuln_detalles): hosts = hosts + " TLS 1.0 habilitado, "
                if ("SSLv3 esta habilitado" in vuln_detalles): hosts = hosts + " SSL 3.0 habilitado, "
                if ("SSLv2 esta habilitado" in vuln_detalles): hosts = hosts + " SSL 2.0 habilitado, "
                if ("HSTS" in vuln_detalles): hosts = hosts + " HSTS deshabilitado, "
                if ("TLS 1.3 no habilitado" in vuln_detalles): hosts = hosts + " TLS 1.3 deshabilitado "
                filas = filas + 1

        if ("vulTLS" in codVul):
            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]
                
                #Convertir salida de herramienta a imagen
                logPath = f'{ruta}logs/vulnerabilidades/{ip}_{port}_{vulType}.txt'
                print (f'logPath vulnTLS {logPath}')
                log2Image(logPath,codVul+'.png')

                vuln_detalles = row[3]
                hosts = hosts + f"\n-{ip}:{port}: {vuln_detalles}"
                filas = filas + 1

        if ("passwordHost" in codVul):
            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]
                
                #Convertir salida de herramienta a imagen
                #log2Image(logPath,codVul+'.png')  

                vuln_detalles = row[3]
                hosts = hosts + f"\n-{vuln_detalles}"
                filas = filas + 1

        ## Mostrar ip puerto y detalles de la vulnerabilidad
        if codVul in ['owaVul', 'cmsDesactualizado']:
            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]
                if codVul not in vuln_no_log:
                #Convertir salida de herramienta a imagen
                    logPath = f'{ruta}logs/vulnerabilidades/{ip}_{port}_{vulType}.txt'
                    log2Image(logPath,codVul+'.png')  

                vuln_detalles = "La " + row[3]
                vuln_detalles = vuln_detalles.replace("VULNERABLE", "es vulnerable")
                hosts = hosts + f" {ip}:{port} : " + vuln_detalles + "\n"
                filas = filas + 1

        if codVul in ['github','scribd']:
            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]
                
                #Convertir salida de herramienta a imagen
                logPath = f'{ruta}logs/enumeracion/{ip}_{port}_{vulType}.txt'
                log2Image(logPath,codVul+'.png')

                vuln_detalles = row[3]
                vuln_detalles = vuln_detalles.replace("Mensaje de error", "")
                vuln_detalles = vuln_detalles.replace("Posible Backdoor", "")
                vuln_detalles = vuln_detalles.replace("200", "")
                vuln_detalles = vuln_detalles.replace("TRACE", "")

                # https://sigec.fonadin.gob.bo:443/.git/	,
                hosts = hosts + ip +":"+ port + "\n" + vuln_detalles + "\n"
                filas = filas + 1


        if codVul in ["passwordAdivinadoWin","passwordAdivinadoServ","passwordDefecto",'passwordAdminWeb','passwordPhpMyadmin']:
            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]
                
                #Convertir salida de herramienta a imagen
                if codVul not in vuln_no_log:
                    if codVul in 'passwordDefecto':
                        logPath = f'{ruta}logs/vulnerabilidades/{ip}_{port}_{vulType}.txt'
                    else:
                        logPath = f'{ruta}logs/cracking/{ip}_{port}_{vulType}.txt'
                    log2Image(logPath,codVul+'.png')  

                vuln_detalles = row[3]
                vuln_detalles = vuln_detalles.replace("Password encontrado:", "")
                vuln_detalles = vuln_detalles.replace("[FTP] ACCOUNT FOUND:", "")
                vuln_detalles = vuln_detalles.replace("ACCOUNT FOUND:", "")
                vuln_detalles = vuln_detalles.replace("[445]", "")
                vuln_detalles = vuln_detalles.replace("[SUCCESS]", "")
                vuln_detalles = vuln_detalles.replace("][", "")

                # [Tomcat] $line (Usuario:tomcat Password:tomcat)
                # [Cisco] Usuario:cisco $respuesta"
                # Password encontrado: [PRTG] $url Usuario:$user Password:$password
                # [AdminWeb] Usuario:admin $respuesta  #401
                # [445][smb] host: 10.0.0.141   login: administrator   password: Pa$$w0rd
                # Password encontrado: [Pentaho] $url (Usuario:$user Password:$password)
                # [FTP] ACCOUNT FOUND: [ftp] Host: 10.0.2.187 User: root Password:  [SUCCESS]
                # ACCOUNT FOUND: [ftp] Host: 10.0.2.187 User: ftp Password:  [SUCCESS]
                # [AdminWeb] Usuario:admin 18:38:58 patator    INFO - 200  12563:-1       0.074 | gainza                             |    51 | HTTP/1.1 200 OK

                hosts = hosts + f" {vuln_detalles}\n"

                # if($vuln_detalles =~ /Tomcat|Pentaho|AdminWeb/i){$servidores++;}
                # if($vuln_detalles =~ /Cisco|PRTG/i){$dispositivosRed++;}
                # if($vuln_detalles =~ /smb/i){$estacionesTrabajo++;}
                filas = filas + 1

        if ("vrfyHabilitado" in codVul):
            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]
                
                #Convertir salida de herramienta a imagen
                logPath = f'{ruta}logs/vulnerabilidades/{ip}_{port}_{vulType}.txt'
                log2Image(logPath,codVul+'.png')  

                hosts = hosts + ip + "\n"
                filas = filas + 1

        if ("compartidoSMB" in codVul):            
            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]   
                #Convertir salida de herramienta a imagen
                if '-' in vulType:   
                    logPath = f'{ruta}logs/enumeracion/{ip}_{port}_{vulType}.txt'
                    print(f'codVul {codVul} vulType {vulType}')
                else:
                    logPath = f'{ruta}logs/vulnerabilidades/{ip}_{port}_{vulType}.txt'

                if vulType != 'sensible-compartidoSMB':   
                    log2Image(logPath,codVul+'.png')  

                vuln_detalles = row[3]
                vuln_detalles = vuln_detalles.replace("                                             	", " ")
                vuln_detalles = vuln_detalles.replace("READ, WRITE", " ")
                vuln_detalles = vuln_detalles.replace("READ ONLY", " ")
                hosts = hosts + f"\\\\{ip} {vuln_detalles} "
                filas = filas + 1

        if ("winboxVuln" in codVul):
            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]
                
                #Convertir salida de herramienta a imagen
                logPath = f'{ruta}logs/vulnerabilidades/{ip}_{port}_{vulType}.txt'
                log2Image(logPath,codVul+'.png')  

                vuln_detalles = row[3]
                vuln_detalles = vuln_detalles.replace(": \n", ":[vacio]")
                vuln_detalles = vuln_detalles.replace("User", "Usuario")
                vuln_detalles = vuln_detalles.replace("Pass", "Contraseña")
                hosts = hosts + f" {ip} (WinBox - MikroTik) - {vuln_detalles} "
                filas = filas + 1

        if ("snmpCommunity" in codVul):
            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]
                
                #Convertir salida de herramienta a imagen
                logPath = f'{ruta}logs/vulnerabilidades/{ip}_{port}_{vulType}.txt'
                log2Image(logPath,codVul+'.png')  

                vuln_detalles = row[3]
                stream = os.popen('echo "' + vuln_detalles + '" | grep --color=never "Community string"')
                community_string = stream.read()
                hosts = hosts + f"{ip}: {community_string} \n"
                filas = filas + 1

        if ("transferenciaDNS" in codVul):
            for row in resSQL:
                ip = row[0]
                port = row[1]
                vulType = row[2]
                
                #Convertir salida de herramienta a imagen
                logPath = f'{ruta}logs/vulnerabilidades/{ip}_{port}_{vulType}.txt'
                log2Image(logPath,codVul+'.png')  

                vuln_detalles = row[3]
                hosts = hosts + f"{ip}:{port} - Archivo de configuración: {vuln_detalles}\n"
                filas = filas + 1

        if filas > 1:
            print(f'recomendacionMatriz {recomendacionMatriz}')
            detalle_activo = create_detalle_activo(ip, port)                 
            ######## RADICAL #######
            if (empresa == "RADICAL"):
                #Cargar los datos de las vulnerabilidades a un array de diccionarios
                globals()['matrizRecomendaciones' + "_" + vectorInforme].append(
                    {'vectorInforme': vectorInforme, 
                    'activo': activo, 
                    "detalle_activo":detalle_activo, 
                    'impacto_negocio': impacto_negocio, 
                    'vulnerabilidad': nombre, 
                    'descripcion': descripcion, 
                    'riesgo': riesgoInforme,
                    'score':float(score),                     
                    "agente_amenaza":agente_amenaza, 
                    "recomendacionMatriz": recomendacionMatriz, 
                    "conclusion": conclusion})                 


                globals()['wb' + empresa + "-" +vectorInforme].create_sheet()
                sheet = globals()['wb' + empresa + "-" +vectorInforme]['Sheet']

                sheet.column_dimensions['A'].width = 23
                sheet.column_dimensions['B'].width = 32
                sheet.column_dimensions['C'].width = 27

                sheet.title = str(numeroVulnerabilidad)
                if ("CRÍTICO" in riesgoInforme):
                    sheet['A1'].fill = CRÍTICOFill
                    sheet['C5'].fill = CRÍTICOFill

                if ("ALTO" in riesgoInforme):
                    sheet['A1'].fill = altoFill
                    sheet['C5'].fill = altoFill

                if ("MEDIO" in riesgoInforme):
                    sheet['A1'].fill = medioFill
                    sheet['C5'].fill = medioFill

                if ("BAJO" in riesgoInforme):
                    sheet['A1'].fill = bajoFill
                    sheet['C5'].fill = bajoFill
                
                if ("INFORMATIVO" in riesgoInforme):
                    sheet['A1'].fill = informativaFill
                    sheet['C5'].fill = informativaFill

                # Titulo
                sheet.merge_cells('B1:C1')
                sheet.row_dimensions[1].height = 59
                sheet['A1'].font = Impact36
                sheet['B1'].font = Impact22
                sheet['A1'].border = thin_border
                sheet['B1'].border = thin_border
                sheet['C1'].border = thin_border

                sheet['A1'].alignment = Alignment(horizontal="center", vertical='center')
                sheet['B1'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

                sheet['A1'] = str(numeroVulnerabilidad)
                sheet['B1'] = str(nombre).upper()

                

                # grafico
                sheet.row_dimensions[2].height = 76.5
                img = openpyxl.drawing.image.Image('/usr/share/lanscanner/image.png')
                #print(f'img {type(img)}')
                sheet.add_image(img, "D2")
                sheet['A2'].border = thin_border
                sheet['B2'].border = thin_border
                sheet['C2'].border = thin_border
                sheet.merge_cells('A2:C2')

                #imagen de la vulnerabilidad 
                print (f'codVul {codVul}')
                if codVul not in vuln_no_log :
                    imgVul = openpyxl.drawing.image.Image(codVul+'.png')
                    sheet.add_image(imgVul, "D4")

                # texto del grafico
                sheet.row_dimensions[3].height = 53.2
                sheet['A3'].font = Arial10
                sheet['B3'].font = Arial10
                sheet['C3'].font = Arial10

                sheet['A3'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['B3'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['C3'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

                sheet['A3'].border = thin_border
                sheet['B3'].border = thin_border
                sheet['C3'].border = thin_border

                sheet['A3'] = str(agente_amenaza)
                sheet['B3'] = str(impacto_tecnico)
                sheet['C3'] = str(impacto_negocio)

                # ANALISIS DE RIESGO LABEL
                sheet.merge_cells('A4:C4')
                sheet.row_dimensions[4].height = 29
                sheet['A4'].font = Arial11Bold

                sheet['A4'].border = thin_border
                sheet['B4'].border = thin_border
                sheet['C4'].border = thin_border

                sheet['A4'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['A4'] = "ANALISIS DE RIESGO"
                sheet['A4'].fill = greyFill

                # Riesgos
                sheet.row_dimensions[5].height = 34.5
                sheet['A5'].font = Arial10
                sheet['B5'].font = Arial10
                sheet['C5'].font = Arial11Bold

                sheet['A5'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['B5'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['C5'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

                sheet['A5'].border = thin_border
                sheet['B5'].border = thin_border
                sheet['C5'].border = thin_border

                sheet['A5'] = "Score: " + str(score)
                sheet['B5'] = "Código: " + str(CVE)
                sheet['C5'] = "RIESGO: " + str(riesgoInforme)

                # vector CVSS
                sheet.row_dimensions[6].height = 20
                sheet.merge_cells('B6:C6')
                sheet['A6'] = "Vector CVSS/4.0:"
                sheet['A6'].border = thin_border
                sheet['A6'].font = Arial11Bold
                sheet['B6'] =  str(vector)
                sheet['A6'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['B6'].alignment = Alignment(horizontal="left", vertical='center', wrap_text=True)

                # descripcion LABEL - data
                descriptionHeight = 30 + 20 * (len(descripcion) // 80)
                sheet.row_dimensions[7].height = descriptionHeight
                sheet['A7'].font = Arial11Bold
                sheet['B7'].font = Calibri10

                sheet['A7'].border = thin_border
                sheet['B7'].border = thin_border
                sheet['C7'].border = thin_border

                sheet['A7'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['A7'] = "Descripción de la vulnerabilidad:"
                sheet.merge_cells('B7:C7')
                sheet['B7'].alignment = Alignment(horizontal="left", vertical='center', wrap_text=True)
                sheet['B7'] = descripcion

                #ATT-Tactic
                sheet['A8'].border = thin_border
                sheet['B8'].border = thin_border
                sheet['C8'].border = thin_border
                sheet.merge_cells('B8:C8')
                sheet['A8'].font = Arial11Bold
                sheet['B8'].font = Arial10
                sheet['A8'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['B8'].alignment = Alignment(horizontal="left", vertical='center', wrap_text=True)                
                sheet['A8'] = "MITRE ATT&CK técnica:"                
                sheet['B8'] = ATT_Tactic

                # PRUEBA DE CONCEPTO LABEL
                sheet.merge_cells('A9:C9')
                sheet.row_dimensions[9].height = 29
                sheet['A9'].font = Arial11Bold
                sheet['A9'].border = thin_border
                sheet['A9'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['A9'] = "PRUEBA DE CONCEPTO"
                sheet['A9'].fill = greyFill


                # HOSTS LABEL - DATA
                hostHeight1 = 20 + 20 * (len(hosts) // 80)
                countNL = hosts.count("SALTOLINEA")
                hostHeight2 = (countNL + 1) * 20
                if hostHeight1 > hostHeight2:
                    hostHeight = hostHeight1
                else:
                    hostHeight = hostHeight2
                sheet.row_dimensions[10].height = hostHeight
                sheet['A10'].font = Arial11Bold
                sheet['B10'].font = Calibri10

                sheet['A10'].border = thin_border
                sheet['B10'].border = thin_border
                sheet['C10'].border = thin_border

                sheet['A10'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['A10'] = "Hosts afectados:"

                sheet.merge_cells('B10:C10')
                sheet['B10'].alignment = Alignment(horizontal="left", vertical='center', wrap_text=True)
                hosts = hosts.replace("SALTOLINEA", "\n")
                #print (f' hosts {hosts}')
                sheet['B10'] = ILLEGAL_CHARACTERS_RE.sub(r'', hosts)  

                #condicionesPrevias                
                sheet.merge_cells('B11:C11')
                sheet['A11'].border = thin_border
                sheet['B11'].border = thin_border
                sheet['A11'].font = Arial11Bold
                sheet['B11'].font = Arial10
                sheet['A11'].alignment = Alignment(horizontal="center", vertical='center')
                sheet['B11'].alignment = Alignment(horizontal="justify", vertical='center', wrap_text=True)                
                sheet['A11'] = "Condiciones previas:"                
                sheet['B11'] = condicionesPrevias

                # PRUEBA DE CONCEPTO DATA
                cellHeight = 20 + 20 * (len(PruebaConcepto) // 80)
                sheet.merge_cells('A12:C12')
                sheet.row_dimensions[12].height = cellHeight
                sheet['A12'].font = Arial10

                sheet['A12'].border = thin_border
                sheet['B12'].border = thin_border
                sheet['C12'].border = thin_border

                sheet['A12'].alignment = Alignment(horizontal="justify", vertical='center', wrap_text=True)
                PruebaConcepto = PruebaConcepto.replace("SALTOLINEA", "\n")
                sheet['A12'] = PruebaConcepto


                # CONTRAMEDIDAS - LABEL
                sheet.merge_cells('A13:C13')
                sheet.row_dimensions[13].height = 29
                sheet['A13'].font = Arial11Bold

                sheet['A13'].border = thin_border
                sheet['B13'].border = thin_border
                sheet['C13'].border = thin_border

                sheet['A13'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['A13'] = "CONTRAMEDIDAS"
                sheet['A13'].fill = greyFill

                # CONTRAMEDIDAS
                sheet.merge_cells('A14:C14')
                cellHeight = 20 + 20 * (len(recomendacion) // 80)

                sheet.row_dimensions[14].height = cellHeight
                sheet['A14'].font = Arial10

                sheet['A14'].border = thin_border
                sheet['B14'].border = thin_border
                sheet['C14'].border = thin_border
                sheet['A14'].alignment = Alignment(horizontal="justify", vertical='center', wrap_text=True)
                recomendacion = recomendacion.replace("SALTOLINEA", "\n")
                sheet['A14'] = recomendacion

                # REFERENCIAS - LABEL
                sheet.merge_cells('A15:C15')
                sheet.row_dimensions[15].height = 29
                sheet['A15'].font = Arial11Bold

                sheet['A15'].border = thin_border
                sheet['B15'].border = thin_border
                sheet['C15'].border = thin_border

                sheet['A15'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['A15'] = "REFERENCIAS"
                sheet['A15'].fill = greyFill


                # REFERENCIAS
                sheet.merge_cells('A16:C16')
                cellHeight = 20 + 20 * (len(referenciaweb) // 80)
                sheet.row_dimensions[16].height = cellHeight
                sheet['A16'].font = Arial10

                sheet['A16'].border = thin_border
                sheet['B16'].border = thin_border
                sheet['C16'].border = thin_border

                sheet['A16'].alignment = Alignment(horizontal="justify", vertical='center', wrap_text=True)
                referenciaweb = referenciaweb.replace("SALTOLINEA", "\n")
                referenciaweb = referenciaweb.replace("TAB", "\t")
                sheet['A16'] = referenciaweb



            ##########  RORIK #########
            if (empresa == "RORIK"):
                #Cargar los datos de las vulnerabilidades a un array de diccionarios
                detalle_activo = create_detalle_activo(ip, port)
                globals()['matrizRecomendaciones' + "_" + vectorInforme].append(
                    {'vectorInforme': vectorInforme, 
                    'activo': activo, 
                    "detalle_activo":detalle_activo, 
                    'impacto_negocio': impacto_negocio, 
                    'vulnerabilidad': nombre, 
                    'descripcion': descripcion, 
                    'riesgo': riesgoInforme,
                    'score':float(score),                     
                    "agente_amenaza":agente_amenaza, 
                    "recomendacionMatriz": recomendacionMatriz, 
                    "conclusion": conclusion})                

                print('########')
                print(globals()['matrizRecomendaciones' + "_" + vectorInforme])                
                #time.sleep(10)  
                print('########\n')

                globals()['wb' + empresa + "-" +vectorInforme].create_sheet()
                sheet = globals()['wb' + empresa + "-" +vectorInforme]['Sheet']

                sheet.column_dimensions['A'].width = 20
                sheet.column_dimensions['B'].width = 70

                sheet.title = str(numeroVulnerabilidad)
                #print (f'riesgoInforme {riesgoInforme}')
                
                if ("CRÍTICO" in riesgoInforme):
                    sheet['A5'].font = Font_MuyAlto_RORIK

                if ("ALTO" in riesgoInforme):
                    sheet['A5'].font = Font_Alto_RORIK

                if ("MEDIO" in riesgoInforme):
                    sheet['A5'].font = Font_Moderado_RORIK

                if ("BAJO" in riesgoInforme):
                    sheet['A5'].font = Font_Bajo_RORIK
                
                if ("INFORMATIVO" in riesgoInforme):
                    sheet['A5'].font = Font_muyBajo_RORIK

                    

                # Infome modelo RORIK
                # Titulo
                sheet['A1'].font = Calibri10Bold
                sheet['A1'] = str(nombre)
                sheet['A1'].border = thin_border
                sheet['B1'].border = thin_border

                # Vulnerabilidad LABEL
                sheet.merge_cells('A2:B2')
                sheet['A2'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['A2'].border = thin_border
                sheet['B2'].border = thin_border
                sheet['A2'] = "VULNERABILIDAD"
                sheet['A2'].fill = brownFill
                sheet['A2'].font = Calibri10Bold

                # Vulnerabilidad
                sheet.merge_cells('A3:B3')
                sheet['A3'].alignment = Alignment(horizontal="justify", vertical='center', wrap_text=True)
                sheet['A3'].font = Calibri10                 
                sheet['A3'].border = thin_border
                sheet['B3'].border = thin_border
                sheet['A3'] = descripcion
                
                #imagen del log
                if codVul not in vuln_no_log:
                    imgVul = openpyxl.drawing.image.Image(codVul+'.png')
                    sheet.add_image(imgVul, "D4")             

                # FACTOR DE RIESGO - LABEL
                sheet.merge_cells('A4:B4')
                sheet['A4'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['A4'].border = thin_border
                sheet['B4'].border = thin_border
                sheet['A4'] = "FACTOR DE RIESGO"
                sheet['A4'].fill = brownFill
                sheet['A4'].font = Calibri10Bold

                # FACTOR DE RIESGO
                sheet['A5'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)                    
                sheet['B5'].alignment = Alignment(horizontal="left", vertical='center', wrap_text=True)
                sheet['A5'].border = thin_border
                sheet['B5'].border = thin_border
                sheet['A5'] = riesgoInforme
                sheet['B5'].font = Calibri10
                sheet['B5'] = "CVSS Base Score: " + str(score)

                # Vector:
                sheet['A6'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['A6'].border = thin_border
                sheet['B6'].border = thin_border
                sheet['A6'] = "Vector:"                
                sheet['A6'].font = Calibri10Bold

                sheet['B6'].alignment = Alignment(horizontal="left", vertical='center', wrap_text=True)
                sheet['B6'].border = thin_border
                sheet['B6'] = vector                
                sheet['B6'].font = Calibri10

                # Modo detección:
                sheet['A7'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['A7'].border = thin_border
                sheet['B7'].border = thin_border
                sheet['A7'] = "Modo detección:"                
                sheet['A7'].font = Calibri10Bold

                sheet['B7'].alignment = Alignment(horizontal="left", vertical='center', wrap_text=True)                
                sheet['B7'].border = thin_border
                sheet['B7'] = "Caja negra"                
                sheet['B7'].font = Calibri10

                # EXPLOTACION - LABEL
                sheet.merge_cells('A8:B8')
                sheet['A8'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['A8'].border = thin_border
                sheet['A8'] = "EXPLOTACIÓN"
                sheet['A8'].fill = brownFill
                sheet['A8'].font = Calibri10Bold

                # EXPLOTACION
                sheet.merge_cells('A9:B9')
                sheet['A9'].alignment = Alignment(horizontal="justify", vertical='center', wrap_text=True)
                sheet['A9'].border = thin_border
                sheet['B9'].border = thin_border
                cellHeight = 20 + 20 * (len(PruebaConcepto) // 80)
                sheet.row_dimensions[9].height = cellHeight
                sheet['A9'].font = Calibri10
                PruebaConcepto = PruebaConcepto.replace("SALTOLINEA", "\n")
                sheet['A9'] = "POSITIVA," + PruebaConcepto

                # HOST - LABEL
                sheet.merge_cells('A10:B10')
                sheet['A10'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['A10'].border = thin_border
                sheet['B10'].border = thin_border
                sheet['A10'] = "HOSTS AFECTADOS"
                sheet['A10'].fill = brownFill
                sheet['A10'].font = Calibri10Bold

                # HOST
                sheet.merge_cells('A11:B11')
                sheet['A11'].alignment = Alignment(horizontal="left", vertical='center', wrap_text=True)
                sheet['A11'].border = thin_border
                sheet['B11'].border = thin_border
                hostHeight1 = 20 + 20 * (len(hosts) // 80)
                countNL = hosts.count("SALTOLINEA")
                hostHeight2 = (countNL + 1) * 20
                if hostHeight1 > hostHeight2:
                    hostHeight = hostHeight1
                else:
                    hostHeight = hostHeight2
                sheet.row_dimensions[11].height = hostHeight
                sheet['A11'].font = Calibri10
                hosts = hosts.replace("SALTOLINEA", "\n")                
                sheet['A11'] = ILLEGAL_CHARACTERS_RE.sub(r'', hosts)

                # CONTRAMEDIDAS - LABEL
                sheet.merge_cells('A12:B12')
                sheet['A12'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['A12'].border = thin_border
                sheet['B12'].border = thin_border
                sheet['A12'] = "CONTRAMEDIDAS"
                sheet['A12'].fill = brownFill
                sheet['A12'].font = Calibri10Bold

                # CONTRAMEDIDAS
                sheet.merge_cells('A13:B13')
                sheet['A13'].alignment = Alignment(horizontal="justify", vertical='center', wrap_text=True)
                sheet['A13'].border = thin_border
                sheet['B13'].border = thin_border
                cellHeight = 20 + 20 * (len(recomendacion) // 80)
                sheet.row_dimensions[13].height = cellHeight
                sheet['A13'].font = Calibri10
                recomendacion = recomendacion.replace("SALTOLINEA", "\n")
                recomendacion = recomendacion.replace("TAB", "\t")
                sheet['A13'] = recomendacion

                 # PARA CONOCER MAS ACERCA - LABEL
                sheet.merge_cells('A14:B14')
                sheet['A14'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
                sheet['A14'].border = thin_border
                sheet['B14'].border = thin_border
                sheet['A14'] = "PARA CONOCER MAS ACERCA DE LA VULNERABILIDAD, CONSULTE:"
                sheet['A14'].fill = brownFill
                sheet['A14'].font = Calibri10Bold

                sheet.merge_cells('A15:B15')
                sheet['A15'].alignment = Alignment(horizontal="justify", vertical='center', wrap_text=True)
                sheet['A15'].font = Calibri10                 
                sheet['A15'].border = thin_border
                sheet['B15'].border = thin_border
                sheet['A15'] = referenciaweb



            numeroVulnerabilidad = numeroVulnerabilidad + 1

    ##### CONCLUSIONES Y RECOMENDACIONES ####  
    if indexDB == todos_resultados_len or indexDB == 1: #1 = interno  todos_resultados_len = Externo        
        print("Generando conclusiones y recomendaciones")                    
        globals()['wb' + empresa + "-" +vectorInforme].create_sheet()
        sheet = globals()['wb' + empresa + "-" +vectorInforme]['Sheet']
        sheet.title = 'Conclusiones y recomendaciones'  # Change title.

        #ordenar conclusiones
        globals()['matrizRecomendaciones' + "_" + vectorInforme] = sorted(globals()['matrizRecomendaciones' + "_" + vectorInforme], key=lambda k: k['score'],reverse=True) 
        
        sheet.column_dimensions['A'].width = 100
        sheet.column_dimensions['B'].width = 100
        #print ("Reporte RORIKT")
        i = 2

        
        # Escribir conclusiones
        sheet['A1'] = "Conclusiones"
        for tupla in globals()['matrizRecomendaciones' + "_" + vectorInforme]:                
            conclusion = tupla["conclusion"]                                
            sheet['A' + str(i)] = conclusion
            i = i + 1
        
        #Observaciones y recomendacinoes
        j = 1 # Contar las observacinoes
        i = len(globals()['matrizRecomendaciones' + "_" + vectorInforme]) + 3 #numero de conclusiones
        for tupla in globals()['matrizRecomendaciones' + "_" + vectorInforme]:                
            conclusion = tupla["conclusion"]
            recomendacionMatriz = tupla["recomendacionMatriz"]
            sheet['A' + str(i + 1)] = "Hallazgo " + str(j) + ": " + conclusion
            sheet['A' + str(i + 2)] = "Recomendación " + str(j) + ": " + recomendacionMatriz
            sheet['A' + str(i + 3)] = "\n"                
            i = i + 3
            j = j + 1 


    ###### ESTADISTICAS POR RIESGO ####
    # crear nueva pestania
    if indexDB == todos_resultados_len or indexDB == 1: #1 = interno  todos_resultados_len = Externo 
        print(f'Generando estadistica por riesgo {vectorInforme}')
        globals()['wb' + empresa + "-" +vectorInforme].create_sheet()
        globals()['sheet' + empresa] = globals()['wb' + empresa + "-" +vectorInforme]['Sheet']
        globals()['sheet' + empresa].title = f'estadisticas'  # Change title

        globals()['sheet' + empresa].column_dimensions['A'].width = 20
        globals()['sheet' + empresa].column_dimensions['B'].width = 25
        globals()['sheet' + empresa]['A1'] = "Valor"

        globals()['sheet' + empresa]['A2'] = "Crítico" #""
        globals()['sheet' + empresa]['A3'] = "Alto"
        globals()['sheet' + empresa]['A4'] = "Medio"
        globals()['sheet' + empresa]['A5'] = "Bajo"
        globals()['sheet' + empresa]['A6'] = "Informativas" #""    
    
        globals()['sheet' + empresa]['A1'].border = thin_border
        globals()['sheet' + empresa]['A2'].border = thin_border
        globals()['sheet' + empresa]['A3'].border = thin_border
        globals()['sheet' + empresa]['A4'].border = thin_border
        globals()['sheet' + empresa]['A5'].border = thin_border
        globals()['sheet' + empresa]['A6'].border = thin_border        

        globals()['sheet' + empresa]['B1'].border = thin_border
        globals()['sheet' + empresa]['B2'].border = thin_border
        globals()['sheet' + empresa]['B3'].border = thin_border
        globals()['sheet' + empresa]['B4'].border = thin_border
        globals()['sheet' + empresa]['B5'].border = thin_border
        globals()['sheet' + empresa]['B6'].border = thin_border

        globals()['sheet' + empresa]['A1'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['A2'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['A3'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['A4'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['A5'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['A6'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

        globals()['sheet' + empresa]['B1'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['B2'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['B3'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['B4'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['B5'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
        globals()['sheet' + empresa]['B6'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

        globals()['sheet' + empresa]['A1'].fill = negroFill
        globals()['sheet' + empresa]['A1'].font = Arial12BoldWhite
        globals()['sheet' + empresa]['A2'].fill = CRÍTICOFill
        globals()['sheet' + empresa]['A3'].fill = AltoFill2
        globals()['sheet' + empresa]['A4'].fill = ModeradoFill
        globals()['sheet' + empresa]['A5'].fill = bajoFill2
        globals()['sheet' + empresa]['A6'].fill = informativaFill

        globals()['sheet' + empresa]['B1'].fill = negroFill
        globals()['sheet' + empresa]['B1'].font = Arial12BoldWhite
        globals()['sheet' + empresa]['B1'] = "Número de riesgos"
        globals()['sheet' + empresa]['B2'] = globals()['total_vul_criticas_' + vectorInforme]
        globals()['sheet' + empresa]['B3'] = globals()['total_vul_altas_' + vectorInforme]
        globals()['sheet' + empresa]['B4'] = globals()['total_vul_medias_' + vectorInforme]
        globals()['sheet' + empresa]['B5'] = globals()['total_vul_bajas_' + vectorInforme]    
        globals()['sheet' + empresa]['B6'] = globals()['total_vul_info_' + vectorInforme]

        chart = PieChart3D()
        chart.title = f'RIESGOS {vectorInforme}S'
        

        # create data for plotting
        labels = Reference( globals()['sheet' + empresa], min_col=1, min_row=2, max_row=6)
        data = Reference( globals()['sheet' + empresa], min_col=2, min_row=2, max_row=6)

        # adding data to the Doughnut chart object
        chart.add_data(data, titles_from_data=False)
        chart.set_categories(labels)

        chart.dataLabels = DataLabelList()
        chart.dataLabels.showPercent = True
        chart.dataLabels.showVal = False
        chart.dataLabels.showLegendKey = False
        chart.dataLabels.showCatName = False

        # set style of the chart
        chart.style = 26

        # try to set color blue (0000FF) for the 2nd wedge (idx=1) in the series
        series = chart.series[0]
        pt = DataPoint(idx=0)
        pt.graphicalProperties.solidFill = "FF0000"
        series.dPt.append(pt)

        pt = DataPoint(idx=1)
        pt.graphicalProperties.solidFill = "FFC000"
        series.dPt.append(pt)

        pt = DataPoint(idx=2)
        pt.graphicalProperties.solidFill = "FFFF00"
        series.dPt.append(pt)

        pt = DataPoint(idx=3)
        pt.graphicalProperties.solidFill = "92D050"
        series.dPt.append(pt)

        pt = DataPoint(idx=4)
        pt.graphicalProperties.solidFill = "95B3D7"
        series.dPt.append(pt)
        
        
        #adicionar la grafica a la hoja de calculo
        globals()['sheet' + empresa].add_chart(chart, 'C5')
    
    print (f'Escribiendo en  archivo {empresa}-{vectorInforme}.xlsx')        
    globals()['wb' + empresa + "-" +vectorInforme].save(f'{empresa}-{vectorInforme}.xlsx')

print ("\n\n")


###### MATRIZ DE VULNERABILIDADES ####
matrizRecomendaciones_EXTERNO = sorted(matrizRecomendaciones_EXTERNO, key=lambda k: k['score'],reverse=True) 
matrizRecomendaciones_INTERNO = sorted(matrizRecomendaciones_INTERNO, key=lambda k: k['score'],reverse=True) 
# Agrupar y ordenar vulnerabilidades
matrizRecomendaciones = matrizRecomendaciones_EXTERNO + matrizRecomendaciones_INTERNO
#print (matrizRecomendaciones)

wbMatrizVul.create_sheet()
sheet = wbMatrizVul['Sheet']
sheet.title = 'Matriz'  # Change title

sheet.column_dimensions['A'].width = 5
sheet.column_dimensions['B'].width = 12
sheet.column_dimensions['C'].width = 15
sheet.column_dimensions['D'].width = 22
sheet.column_dimensions['E'].width = 25
sheet.column_dimensions['F'].width = 15
sheet.column_dimensions['G'].width = 32
sheet.column_dimensions['H'].width = 15
sheet.column_dimensions['I'].width = 32



sheet['D1'] = "MATRIZ DE VULNERABILIDADES"
sheet['D1'].font = Calibri22Bold

sheet['A2'] = "Clasificación: Confidencial"
sheet['A3'] = "Fecha: Al "
sheet['A2'].font = Calibri12Bold
sheet['A3'].font = Calibri12Bold


# Titulos
sheet['A5'] = "Nº"
sheet['B5'] = "Vector"
sheet['C5'] = "Activo"
sheet['D5'] = "Detalle Activo de información"
sheet['E5'] = "Amenaza"
sheet['F5'] = "Vulnerabilidad identificada"
sheet['G5'] = "Descripcion"
sheet['H5'] = "Riesgo"
sheet['I5'] = "Recomendación"

sheet['A5'].font = Arial14Bold
sheet['B5'].font = Arial14Bold
sheet['C5'].font = Arial14Bold
sheet['D5'].font = Arial14Bold
sheet['E5'].font = Arial14Bold
sheet['F5'].font = Arial14Bold
sheet['G5'].font = Arial14Bold
sheet['H5'].font = Arial14Bold
sheet['I5'].font = Arial14Bold
sheet['J5'].font = Arial14Bold

sheet['A5'].alignment = Alignment(horizontal="center", vertical='center',wrap_text=True)
sheet['B5'].alignment = Alignment(horizontal="center", vertical='center',wrap_text=True)
sheet['C5'].alignment = Alignment(horizontal="center", vertical='center',wrap_text=True)
sheet['D5'].alignment = Alignment(horizontal="center", vertical='center',wrap_text=True)
sheet['E5'].alignment = Alignment(horizontal="center", vertical='center',wrap_text=True)
sheet['F5'].alignment = Alignment(horizontal="center", vertical='center',wrap_text=True)
sheet['G5'].alignment = Alignment(horizontal="center", vertical='center',wrap_text=True)
sheet['H5'].alignment = Alignment(horizontal="center", vertical='center',wrap_text=True)
sheet['I5'].alignment = Alignment(horizontal="center", vertical='center',wrap_text=True)
sheet['J5'].alignment = Alignment(horizontal="center", vertical='center',wrap_text=True)

sheet['A5'].border = thin_border
sheet['B5'].border = thin_border
sheet['C5'].border = thin_border
sheet['D5'].border = thin_border
sheet['E5'].border = thin_border
sheet['F5'].border = thin_border
sheet['G5'].border = thin_border
sheet['H5'].border = thin_border
sheet['I5'].border = thin_border
sheet['J5'].border = thin_border

# Vulnerabilidades
i = 1
#print (matrizRecomendaciones)
for tupla in matrizRecomendaciones:
    #print (vulnerabilidad)
    vectorInforme = tupla["vectorInforme"]   
    activo = tupla["activo"]
    detalle_activo = tupla["detalle_activo"]
    amenaza = tupla["impacto_negocio"]
    vulnerabilidad = tupla["vulnerabilidad"]
    descripcion_vulnerabilidad = tupla["descripcion"]
    riesgo = tupla['riesgo']
    score = float(tupla["score"])
    recomendacionMatriz = tupla["recomendacionMatriz"]      
    

    sheet['A' + str(i+5)] = i
    sheet['B' + str(i+5)] = vectorInforme
    sheet['C' + str(i+5)] = activo
    sheet['D' + str(i+5)] = detalle_activo
    sheet['E' + str(i+5)] = amenaza
    sheet['F' + str(i+5)] = vulnerabilidad
    sheet['G' + str(i+5)] = descripcion_vulnerabilidad
    sheet['H' + str(i+5)] = riesgo
    sheet['I' + str(i+5)] = recomendacionMatriz
    sheet['J' + str(i+5)] = score



    sheet['A' + str(i + 5)].alignment = Alignment(horizontal="center", vertical='center',wrap_text=True)
    sheet['B' + str(i + 5)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    sheet['C' + str(i + 5)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    sheet['D' + str(i + 5)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    sheet['E' + str(i + 5)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    sheet['F' + str(i + 5)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    sheet['G' + str(i + 5)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    sheet['H' + str(i + 5)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    sheet['I' + str(i + 5)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
    sheet['J' + str(i + 5)].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

    sheet['A' + str(i + 5)].border = thin_border
    sheet['B' + str(i + 5)].border = thin_border
    sheet['C' + str(i + 5)].border = thin_border
    sheet['D' + str(i + 5)].border = thin_border
    sheet['E' + str(i + 5)].border = thin_border
    sheet['F' + str(i + 5)].border = thin_border
    sheet['G' + str(i + 5)].border = thin_border
    sheet['H' + str(i + 5)].border = thin_border
    sheet['I' + str(i + 5)].border = thin_border
    sheet['J' + str(i + 5)].border = thin_border

    if ("CRÍTICO" in riesgo):
        sheet['H' + str(i+5)].fill = CRÍTICOFill

    if ("ALTO" in riesgo):
        sheet['H' + str(i+5)].fill = altoFill

    if ("MEDIO" in riesgo):
        sheet['H' + str(i+5)].fill = medioFill
    
    if ("BAJO" in riesgo):
        sheet['H' + str(i+5)].fill = bajoFill

    if ("INFORMATIVO" in riesgo):
        sheet['H' + str(i+5)].fill = informativaFill
    i = i + 1


wbMatrizVul.save(f'Matriz-vulnerabilidades.xlsx')    

###### INFORME EJECUTIVO ####
wbEjecutivo.create_sheet()
sheet = wbEjecutivo['Sheet']

sheet.title = 'Vulnerabilidades identificadas'  # Change title

total_vul_criticas = total_vul_criticas_INTERNO + total_vul_criticas_EXTERNO
total_vul_altas = total_vul_altas_INTERNO + total_vul_altas_EXTERNO
total_vul_medias = total_vul_medias_INTERNO + total_vul_medias_EXTERNO
total_vul_bajas = total_vul_bajas_INTERNO + total_vul_bajas_EXTERNO
total_vul_info= total_vul_info_INTERNO + total_vul_info_EXTERNO

vulnerabilidades_texto_ejecutivo = generar_texto_vulnerabilidades(total_vul_criticas, total_vul_altas, total_vul_medias, total_vul_bajas, total_vul_info)

############## total vulnerabilidades #################
sheet.column_dimensions['A'].width = 20
sheet.column_dimensions['B'].width = 25
sheet['A1'] = "Valor"
sheet['A2'] = "Crítico"
sheet['A3'] = "Alto"
sheet['A4'] = "Medio"
sheet['A5'] = "Bajo"
sheet['A6'] = "Informativas"    
sheet['C1'] = vulnerabilidades_texto_ejecutivo 

sheet['A1'].border = thin_border
sheet['A2'].border = thin_border
sheet['A3'].border = thin_border
sheet['A4'].border = thin_border
sheet['A5'].border = thin_border
sheet['A6'].border = thin_border

sheet['B1'].border = thin_border
sheet['B2'].border = thin_border
sheet['B3'].border = thin_border
sheet['B4'].border = thin_border
sheet['B5'].border = thin_border
sheet['B6'].border = thin_border

sheet['A1'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A2'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A3'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A4'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A5'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A6'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

sheet['B1'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B2'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B3'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B4'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B5'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B6'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

sheet['A1'].fill = negroFill
sheet['A1'].font = Arial12BoldWhite
sheet['A2'].fill = CRÍTICOFill
sheet['A3'].fill = AltoFill2
sheet['A4'].fill = ModeradoFill   
sheet['A5'].fill = bajoFill2
sheet['A6'].fill = informativaFill

sheet['B1'].fill = negroFill
sheet['B1'].font = Arial12BoldWhite
sheet['B1'] = "Número de riesgos"
sheet['B2'] = total_vul_criticas
sheet['B3'] = total_vul_altas
sheet['B4'] = total_vul_medias
sheet['B5'] = total_vul_bajas
sheet['B6'] = total_vul_info
#################################


############# Vul externas ##########
sheet['A9'] = "Sumario de Hallazgos externos"
sheet['A10']= "Valor"
sheet['A11'] = "Crítico"
sheet['A12'] = "Alto"
sheet['A13'] = "Medio"
sheet['A14'] = "Bajo"
sheet['A15'] = "Informativas"    

sheet['A10'].border = thin_border
sheet['A11'].border = thin_border
sheet['A12'].border = thin_border
sheet['A13'].border = thin_border
sheet['A14'].border = thin_border
sheet['A15'].border = thin_border

sheet['B10'].border = thin_border
sheet['B11'].border = thin_border
sheet['B12'].border = thin_border
sheet['B13'].border = thin_border
sheet['B14'].border = thin_border
sheet['B15'].border = thin_border

sheet['A10'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A11'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A12'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A13'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A14'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A15'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

sheet['B10'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B11'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B12'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B13'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B14'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B15'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

sheet['A10'].fill = negroFill
sheet['A10'].font = Arial12BoldWhite
sheet['A11'].fill = CRÍTICOFill
sheet['A12'].fill = AltoFill2
sheet['A13'].fill = ModeradoFill   
sheet['A14'].fill = bajoFill2
sheet['A15'].fill = informativaFill

sheet['B10'].fill = negroFill
sheet['B10'].font = Arial12BoldWhite
sheet['B10'] = "Número de riesgos"
sheet['B11'] = total_vul_criticas_EXTERNO
sheet['B12'] = total_vul_altas_EXTERNO
sheet['B13'] = total_vul_medias_EXTERNO
sheet['B14'] = total_vul_bajas_EXTERNO
sheet['B15'] = total_vul_info_EXTERNO
#################################


############# Vul internas ##########
sheet['A17'] = "Sumario de Hallazgos internos"
sheet['A18']= "Valor"
sheet['A19'] = "Crítico"
sheet['A20'] = "Alto"
sheet['A21'] = "Medio"
sheet['A22'] = "Bajo"
sheet['A23'] = "Informativas"    

sheet['A18'].border = thin_border
sheet['A19'].border = thin_border
sheet['A20'].border = thin_border
sheet['A21'].border = thin_border
sheet['A22'].border = thin_border
sheet['A23'].border = thin_border

sheet['B18'].border = thin_border
sheet['B19'].border = thin_border
sheet['B20'].border = thin_border
sheet['B21'].border = thin_border
sheet['B22'].border = thin_border
sheet['B23'].border = thin_border

sheet['A18'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A19'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A20'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A21'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A22'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A23'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

sheet['B18'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B19'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B20'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B21'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B22'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B23'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

sheet['A18'].fill = negroFill
sheet['A18'].font = Arial12BoldWhite
sheet['A19'].fill = CRÍTICOFill
sheet['A20'].fill = AltoFill2
sheet['A21'].fill = ModeradoFill   
sheet['A22'].fill = bajoFill2
sheet['A23'].fill = informativaFill

sheet['B18'].fill = negroFill
sheet['B18'].font = Arial12BoldWhite
sheet['B18'] = "Número de riesgos"
sheet['B19'] = total_vul_criticas_INTERNO
sheet['B20'] = total_vul_altas_INTERNO
sheet['B21'] = total_vul_medias_INTERNO
sheet['B22'] = total_vul_bajas_INTERNO
sheet['B23'] = total_vul_info_INTERNO
#################################

chart = BarChart()
chart.type = "col"
chart.style = 10
chart.shape = 4
chart.title = "Vulnerabilidades identificadas"

# create data for plotting
labels = Reference( sheet, min_col=1, min_row=2, max_row=6)
data = Reference( sheet, min_col=2, min_row=2, max_row=6)

# adding data to the Doughnut chart object
chart.add_data(data, titles_from_data=False)
chart.set_categories(labels)

chart.dataLabels = DataLabelList()
chart.dataLabels.showPercent = True
chart.dataLabels.showVal = False
chart.dataLabels.showLegendKey = True
chart.dataLabels.showCatName = False

# try to set color blue (0000FF) for the 2nd wedge (idx=1) in the series
series = chart.series[0]
pt = DataPoint(idx=0)
pt.graphicalProperties.solidFill = "FF0000"
series.dPt.append(pt)

pt = DataPoint(idx=1)
pt.graphicalProperties.solidFill = "FFC000"
series.dPt.append(pt)

pt = DataPoint(idx=2)
pt.graphicalProperties.solidFill = "FFFF00"
series.dPt.append(pt)

pt = DataPoint(idx=3)
pt.graphicalProperties.solidFill = "92D050"
series.dPt.append(pt)

pt = DataPoint(idx=4)
pt.graphicalProperties.solidFill = "95B3D7"
series.dPt.append(pt)

#adicionar la grafica a la hoja de calculo
sheet.add_chart(chart, 'C2')




###### Vulnerabilidades externas vs internas ####
        # crear nueva pestania
wbEjecutivo.create_sheet()
sheet = wbEjecutivo['Sheet']

sheet.title = 'externas vs internas'  # Change title

sheet.column_dimensions['A'].width = 20
sheet.column_dimensions['B'].width = 25
sheet['A1'] = "Vector"
sheet['A2'] = "Externas"
sheet['A3'] = "Internas"


sheet['A1'].border = thin_border
sheet['A2'].border = thin_border
sheet['A3'].border = thin_border


sheet['B1'].border = thin_border
sheet['B2'].border = thin_border
sheet['B3'].border = thin_border


sheet['A1'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A2'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A3'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)


sheet['B1'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B2'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B3'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)



sheet['A1'].fill = negroFill
sheet['A1'].font = Arial12BoldWhite
sheet['A1'] = "Vector"
sheet['B1'].fill = negroFill
sheet['B1'].font = Arial12BoldWhite
sheet['B1'] = "Vulnerabilidades"

vulnerabilidades_externas = total_vul_criticas_EXTERNO + total_vul_altas_EXTERNO + total_vul_medias_EXTERNO + total_vul_bajas_EXTERNO + total_vul_info_EXTERNO
vulnerabilidades_internas = total_vul_criticas_INTERNO + total_vul_altas_INTERNO +total_vul_medias_INTERNO + total_vul_bajas_INTERNO + total_vul_info_INTERNO
total_vulnerabilidades = vulnerabilidades_externas + vulnerabilidades_internas
sheet['B2'] = vulnerabilidades_externas
sheet['B3'] = vulnerabilidades_internas
print (f'vulnerabilidades_externas {vulnerabilidades_externas}')
print (f'vulnerabilidades_internas {vulnerabilidades_internas}')

chart = PieChart3D()
chart.title = "Vulnerabilidades por vector"
# create data for plotting
labels = Reference( sheet, min_col=1, min_row=2, max_row=3)
data = Reference( sheet, min_col=2, min_row=2, max_row=3)

# adding data to the Doughnut chart object
chart.add_data(data, titles_from_data=False)
chart.set_categories(labels)

chart.dataLabels = DataLabelList()
chart.dataLabels.showPercent = True
chart.dataLabels.showVal = False
chart.dataLabels.showLegendKey = False
chart.dataLabels.showCatName = False

# set style of the chart
chart.style = 26

# try to set color blue (0000FF) for the 2nd wedge (idx=1) in the series
series = chart.series[0]
pt = DataPoint(idx=0)
pt.graphicalProperties.solidFill = "B74C49"
series.dPt.append(pt)

pt = DataPoint(idx=1)
pt.graphicalProperties.solidFill = "4B7BB4"
series.dPt.append(pt)

#adicionar la grafica a la hoja de calculo
sheet.add_chart(chart, 'C5')



###### Vulnerabilidades por tipo de activo ####
        # crear nueva pestania
wbEjecutivo.create_sheet()
sheet = wbEjecutivo['Sheet']

sheet.title = 'activos de información'  # Change title

sheet.column_dimensions['A'].width = 20
sheet.column_dimensions['B'].width = 25

sheet['A2'] = "Aplicación de escritorio"
sheet['A3'] = "Aplicación móvil"
sheet['A4'] = "Aplicación web"
sheet['A5'] = "ATM"
sheet['A6'] = "Estación de trabajo"
sheet['A7'] = "Red WiFi"
sheet['A8'] = "Servidor"
sheet['A9'] = "Base de datos"
sheet['A10'] = "VoIP"
sheet['A11'] = "Sistemas de vigilancia"
sheet['A12'] = "Dispositivos de red"
sheet['A13'] = "Personal"
sheet['A14'] = "Otros"


sheet['A1'].border = thin_border
sheet['A2'].border = thin_border
sheet['A3'].border = thin_border
sheet['A4'].border = thin_border
sheet['A5'].border = thin_border
sheet['A6'].border = thin_border
sheet['A7'].border = thin_border
sheet['A8'].border = thin_border
sheet['A9'].border = thin_border
sheet['A10'].border = thin_border
sheet['A11'].border = thin_border
sheet['A12'].border = thin_border
sheet['A13'].border = thin_border
sheet['A14'].border = thin_border



sheet['B1'].border = thin_border
sheet['B2'].border = thin_border
sheet['B3'].border = thin_border
sheet['B4'].border = thin_border
sheet['B5'].border = thin_border
sheet['B6'].border = thin_border
sheet['B7'].border = thin_border
sheet['B8'].border = thin_border
sheet['B9'].border = thin_border
sheet['B10'].border = thin_border
sheet['B11'].border = thin_border
sheet['B12'].border = thin_border
sheet['B13'].border = thin_border
sheet['B14'].border = thin_border


sheet['A1'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A2'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A3'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A4'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A5'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A6'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A7'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A8'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A9'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A10'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A11'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A12'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A13'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A14'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)


sheet['B1'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B2'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B3'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B4'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B5'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B6'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B7'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B8'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B9'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B10'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B11'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B12'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B13'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B14'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)

## Cabecera
sheet['A1'].fill = negroFill
sheet['A1'].font = Arial12BoldWhite
sheet['A1'] = "Tipo de activo"
sheet['B1'].fill = negroFill
sheet['B1'].font = Arial12BoldWhite
sheet['B1'] = "Vulnerabilidades"

#datos
sheet['B2'] = aplicacionEscritorio
sheet['B3'] = aplicacionMovil
sheet['B4'] = aplicacionWeb
sheet['B5'] = atm
sheet['B6'] = estacionesTrabajo
sheet['B7'] = wifi
sheet['B8'] = servidores
sheet['B9'] = baseDatos
sheet['B10'] = telefoniaIP
sheet['B11'] = sistemaVigilancia
sheet['B12'] = dispositivosRed
sheet['B13'] = personal
sheet['B14'] = otros


chart = PieChart()
chart.title = "Vulnerabilidades por activos de información"
#chart.style = 26

# create data for plotting
labels = Reference( sheet, min_col=1, min_row=2, max_row=10)
data = Reference( sheet, min_col=2, min_row=2, max_row=10)

# adding data to the Doughnut chart object
chart.add_data(data, titles_from_data=False)
chart.set_categories(labels)

# agregar espacio
slice = DataPoint(idx=0, explosion=20)
chart.series[0].data_points = [slice]
#chart.series[1].data_points = [slice]
#chart.series[2].data_points = [slice]

chart.dataLabels = DataLabelList()
chart.dataLabels.showPercent = True
chart.dataLabels.showVal = False
chart.dataLabels.showLegendKey = True
chart.dataLabels.showCatName = True

#adicionar la grafica a la hoja de calculo
sheet.add_chart(chart, 'C5')



### Total pruebas vs explotadas

# crear nueva pestania
wbEjecutivo.create_sheet()
sheet = wbEjecutivo['Sheet']

sheet.title = 'Total pruebas'  # Change title

sheet.column_dimensions['A'].width = 20
sheet.column_dimensions['B'].width = 25
sheet['A1'] = "Total pruebas específicas"
sheet['A2'] = "Vulnerabilidades identificadas"

sheet['A1'].border = thin_border
sheet['A2'].border = thin_border


sheet['B1'].border = thin_border
sheet['B2'].border = thin_border

sheet['A1'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A2'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)


sheet['B1'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B2'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)



sheet['B1'] = totalPruebas
sheet['B2'] = total_vul_criticas_INTERNO + total_vul_criticas_EXTERNO + total_vul_altas_INTERNO + total_vul_altas_EXTERNO +total_vul_medias_INTERNO + total_vul_medias_EXTERNO + total_vul_bajas_INTERNO + total_vul_bajas_EXTERNO + total_vul_info_EXTERNO + total_vul_info_INTERNO



chart = BarChart()
chart.type = "bar"
chart.style = 10
chart.shape = 4
chart.title = "Vulnerabilidades identificadas"

# create data for plotting
labels = Reference( sheet, min_col=1, min_row=1, max_row=2)
data = Reference( sheet, min_col=2, min_row=1, max_row=2)

# adding data to the Doughnut chart object
chart.add_data(data, titles_from_data=False)
chart.set_categories(labels)

chart.dataLabels = DataLabelList()
chart.dataLabels.showPercent = True
chart.dataLabels.showVal = True
chart.dataLabels.showLegendKey = True
chart.dataLabels.showCatName = False

# try to set color blue (0000FF) for the 2nd wedge (idx=1) in the series
series = chart.series[0]
pt = DataPoint(idx=0)
pt.graphicalProperties.solidFill = "007BD4"
series.dPt.append(pt)

pt = DataPoint(idx=1)
pt.graphicalProperties.solidFill = "FB770B"
series.dPt.append(pt)

#adicionar la grafica a la hoja de calculo
sheet.add_chart(chart, 'C5')




###### Vulnerabilidades por tipo de vulnerabilidad ####
        # crear nueva pestania
wbEjecutivo.create_sheet()
sheet = wbEjecutivo['Sheet']

sheet.title = 'Tipo vulnerabilidad'  # Change title

sheet.column_dimensions['A'].width = 20
sheet.column_dimensions['B'].width = 25

sheet['A2'] = "Password débil"
sheet['A3'] = "Falta parches de seguridad"
sheet['A4'] = "Errores de configuración"
sheet['A5'] = "Programación insegura"

sheet['A1'].border = thin_border
sheet['A2'].border = thin_border
sheet['A3'].border = thin_border
sheet['A4'].border = thin_border
sheet['A5'].border = thin_border


sheet['B1'].border = thin_border
sheet['B2'].border = thin_border
sheet['B3'].border = thin_border
sheet['B4'].border = thin_border
sheet['B5'].border = thin_border


sheet['A1'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A2'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A3'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A4'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['A5'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)


sheet['B1'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B2'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B3'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B4'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)
sheet['B5'].alignment = Alignment(horizontal="center", vertical='center', wrap_text=True)


## Cabecera
sheet['A1'].fill = negroFill
sheet['A1'].font = Arial12BoldWhite
sheet['A1'] = "Tipo"
sheet['B1'].fill = negroFill
sheet['B1'].font = Arial12BoldWhite
sheet['B1'] = "Vulnerabilidades"

#datos
sheet['B2'] = passwordDebil
sheet['B3'] = faltaParches
sheet['B4'] = errorConfiguracion
sheet['B5'] = programacionInsegura

chart = PieChart3D()
chart.title = "Vulnerabilidades por tipo"
#chart.style = 26

# create data for plotting
labels = Reference( sheet, min_col=1, min_row=2, max_row=5)
data = Reference( sheet, min_col=2, min_row=2, max_row=5)

# adding data to the Doughnut chart object
chart.add_data(data, titles_from_data=False)
chart.set_categories(labels)


chart.dataLabels = DataLabelList()
chart.dataLabels.showPercent = True
chart.dataLabels.showVal = False
chart.dataLabels.showLegendKey = False
chart.dataLabels.showCatName = False
#adicionar la grafica a la hoja de calculo
sheet.add_chart(chart, 'C5')
##################################################

# crear nueva pestania
wbEjecutivo.create_sheet()
sheet = wbEjecutivo['Sheet']

sheet.title = 'Conclusiones y recomendaciones'  # Change title
TOTAL_CONCLUSIONES_RECOMENDACIONES = matrizRecomendaciones_INTERNO + matrizRecomendaciones_EXTERNO
sheet['A1'] = "Conclusiones"
sheet['B1'] = "Recomendaciones"
i = 2

for tupla in TOTAL_CONCLUSIONES_RECOMENDACIONES:
    conclusion = tupla["conclusion"]
    recomendacionMatriz = tupla["recomendacionMatriz"] 
    sheet['A' + str(i)] = conclusion 
    sheet['B' + str(i)] = recomendacionMatriz                   
    i = i + 1


wbEjecutivo.save(f'informe-ejecutivo.xlsx')


############# retest ###########
wordRetest = Document()

wordRetest.add_heading("INTRODUCCIÓN", level=1)
paragraph = wordRetest.add_paragraph("Este documento es el resultado de la verificación de la implementación de contramedidas sugeridas en el informe remitido en XXX")
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph.add_run().add_break() 

wordRetest.add_heading("VALIDACIÓN DE CONTRAMEDIDAS EXTERNAS", level=1)
i = 1
for tupla in matrizRecomendaciones_EXTERNO:
    title_vuln = tupla["vulnerabilidad"]
    wordRetest.add_heading(title_vuln, level=2)
    create_custom_table(wordRetest, title_vuln, i)
    wordRetest.add_page_break()  # Add a page break after each table except the last one
    i = i + 1


wordRetest.add_heading("VALIDACIÓN DE CONTRAMEDIDAS INTERNAS", level=1)
i = 1
for tupla in matrizRecomendaciones_INTERNO:
    title_vuln = tupla["vulnerabilidad"]
    wordRetest.add_heading(title_vuln, level=2)
    create_custom_table(wordRetest, title_vuln, i)
    wordRetest.add_page_break()  # Add a page break after each table except the last one
    i = i + 1

wordRetest.save('retest.docx')
########################

############# retest ejecutivo ###########
wordRetestEjec = Document()

wordRetestEjec.add_heading("Introducción", level=1)
paragraph = wordRetestEjec.add_paragraph("Este documento es el resultado de la verificación de la implementación de contramedidas sugeridas en el informe remitido en XXX")
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph.add_run().add_break() 

wordRetestEjec.add_heading("Sumario de Hallazgos", level=1)
paragraph = wordRetestEjec.add_paragraph(f"A continuación, se presentan el estado de las {total_vulnerabilidades} vulnerabilidades identificadas inicialmente")
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph.add_run().add_break() 

# Definir los encabezados de la tabla
tabla = wordRetestEjec.add_table(rows=1, cols=6)
tabla.style = 'Table Grid'
encabezados = ["Nº", "Vector", "Activo", "Vulnerabilidad", "Estado", "Nivel criticidad"]

for i, encabezado in enumerate(encabezados):
    tabla.cell(0, i).text = encabezado
    # Opcional: agregar estilo al texto del encabezado
    run = tabla.cell(0, i).paragraphs[0].runs[0]
    run.font.bold = True
    run.font.size = Pt(12)

# Rellenar las filas de la tabla con los datos
i = 1
for tupla in matrizRecomendaciones:
    row_cells = tabla.add_row().cells
    
    row_cells[0].text = str(i)
    row_cells[1].text = tupla["vectorInforme"]
    row_cells[2].text = tupla["activo"]
    row_cells[3].text = tupla["vulnerabilidad"]
    row_cells[4].text = "Pendiente"  # Estado inicial
    row_cells[5].text = tupla['riesgo']

    i = i +1

    # Aplicar color basado en el nivel de criticidad
    criticidad_color = nivel_criticidad_colores.get(tupla['riesgo'].strip(), "FFFFFF")
    set_cell_background(row_cells[5], criticidad_color)

wordRetestEjec.add_heading("GRÁFICOS ESTADÍSTICOS", level=1)
wordRetestEjec.add_heading("Estado de las vulnerabilidades", level=2)
paragraph = wordRetestEjec.add_paragraph(f"Sumario del estado de las vulnerabilidades")
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph.add_run().add_break() 

wordRetestEjec.add_heading("Vulnerabilidades por nivel de riesgo", level=2)
paragraph = wordRetestEjec.add_paragraph(f"El retest ha evidenciado que aún existen XXX vulnerabilidades con los siguientes niveles de riesgo:")
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph.add_run().add_break() 


wordRetestEjec.save('retestEjec.docx')
########################  
